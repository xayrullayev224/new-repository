import telebot
import sqlite3
import sys
sys.stdout.reconfigure(encoding='utf-8')
import time
import os
import threading
import urllib.request
from datetime import datetime
from google import genai
from groq import Groq
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
from gtts import gTTS
from elevenlabs.client import ElevenLabs
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
from reportlab.graphics.shapes import Drawing
from reportlab.graphics.charts.barcharts import VerticalBarChart
from reportlab.graphics.charts.linecharts import HorizontalLineChart
from reportlab.graphics.widgetbase import Widget
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

# ================== SOZLAMALAR ==================
BOT_TOKEN   = 
GEMINI_KEY  = 
GROQ_KEY    = 
ELEVEN_KEY  = 
ADMIN_ID    = 
WEBAPP_URL  = os.environ.get("WEBAPP_URL", "https://www.sesame.com")
# Mock test uchun vebsayt (O'zingiz netlify.com yoki bepul serverga html faylni yuklab linkini bu yerga yozasiz)
MOCK_WEBAPP_URL  = os.environ.get("MOCK_WEBAPP_URL", "https://xayrullayev224.github.io/mock-test/") # GitHub Pages linki (Bannersiz va Fullscreen uchun)

gemini = None
if GEMINI_KEY:
    try:
        gemini = genai.Client(api_key=GEMINI_KEY)
    except Exception as e:
        print(f"⚠️ Gemini klientini yaratib bo'lmadi: {e}")
else:
    print("⚠️ GEMINI_KEY topilmadi.")

groq = Groq(api_key=GROQ_KEY) if GROQ_KEY else None
eleven = ElevenLabs(api_key=ELEVEN_KEY) if ELEVEN_KEY else None
bot = telebot.TeleBot(BOT_TOKEN, threaded=True)

telebot.apihelper.CONNECT_TIMEOUT = 30
telebot.apihelper.READ_TIMEOUT = 30

waiting_users = {}
active_pairs = {}
user_states = {}
state_lock = threading.Lock()


# ================== ENGNOVATE LINKS ==================

def eng_listening_url(n, t):
    return f"https://engnovate.com/ielts-listening-tests/cambridge-ielts-{n}-academic-listening-test-{t}/"


def eng_reading_url(n, t):
    return f"https://engnovate.com/ielts-reading-tests/cambridge-ielts-{n}-academic-reading-test-{t}/"


ENGNOVATE = {
    "Listening": {
        f"Cambridge {n}": {f"Test {t}": eng_listening_url(n, t) for t in range(1, 5)}
        for n in range(1, 21)
    },
    "Reading": {
        f"Cambridge {n}": {f"Test {t}": eng_reading_url(n, t) for t in range(1, 5)}
        for n in range(1, 21)
    }
}


# ================== DATABASE ==================

def db():
    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    db_path = os.path.join(BASE_DIR, "bot.db")
    conn = sqlite3.connect(db_path, check_same_thread=False)
    conn.row_factory = sqlite3.Row
    return conn


def init_db():
    with db() as conn:
        conn.executescript("""
                           CREATE TABLE IF NOT EXISTS users
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY,
                               telegram_id
                               INTEGER
                               UNIQUE,
                               username
                               TEXT,
                               full_name
                               TEXT,
                               is_premium
                               INTEGER
                               DEFAULT
                               0,
                               joined_date
                               TEXT
                               DEFAULT
                               CURRENT_TIMESTAMP
                           );
                           CREATE TABLE IF NOT EXISTS progress
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY,
                               telegram_id
                               INTEGER
                               UNIQUE,
                               total_questions
                               INTEGER
                               DEFAULT
                               0,
                               correct_answers
                               INTEGER
                               DEFAULT
                               0
                           );
                           CREATE TABLE IF NOT EXISTS materials
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY
                               AUTOINCREMENT,
                               section
                               TEXT,
                               category
                               TEXT,
                               book
                               TEXT,
                               test
                               TEXT,
                               part
                               TEXT,
                               title
                               TEXT,
                               file_id
                               TEXT,
                               file_type
                               TEXT
                               DEFAULT
                               'document',
                               date
                               TEXT
                               DEFAULT
                               CURRENT_TIMESTAMP
                           );
                           CREATE TABLE IF NOT EXISTS history
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY
                               AUTOINCREMENT,
                               telegram_id
                               INTEGER,
                               section
                               TEXT,
                               question
                               TEXT,
                               answer
                               TEXT,
                               date
                               TEXT
                               DEFAULT
                               CURRENT_TIMESTAMP
                           );
                           CREATE TABLE IF NOT EXISTS daily_tasks
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY
                               AUTOINCREMENT,
                               telegram_id
                               INTEGER,
                               task_text
                               TEXT,
                               is_done
                               INTEGER
                               DEFAULT
                               0,
                               date
                               TEXT
                               DEFAULT
                               CURRENT_TIMESTAMP
                           );
                           CREATE TABLE IF NOT EXISTS leaderboard
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY,
                               telegram_id
                               INTEGER
                               UNIQUE,
                               weekly_score
                               INTEGER
                               DEFAULT
                               0,
                               total_score
                               INTEGER
                               DEFAULT
                               0,
                               last_updated
                               TEXT
                               DEFAULT
                               CURRENT_TIMESTAMP
                           );
                           CREATE TABLE IF NOT EXISTS reminders
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY,
                               telegram_id
                               INTEGER
                               UNIQUE,
                               reminder_time
                               TEXT
                               DEFAULT
                               '09:00',
                               is_active
                               INTEGER
                               DEFAULT
                               1
                           );
                           CREATE TABLE IF NOT EXISTS study_plans
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY
                               AUTOINCREMENT,
                               telegram_id
                               INTEGER,
                               current_level
                               REAL,
                               target_level
                               REAL,
                               duration_weeks
                               INTEGER,
                               plan_content
                               TEXT,
                               created_date
                               TEXT
                               DEFAULT
                               CURRENT_TIMESTAMP
                           );
                           CREATE TABLE IF NOT EXISTS weekly_progress
                           (
                               id
                               INTEGER
                               PRIMARY
                               KEY
                               AUTOINCREMENT,
                               telegram_id
                               INTEGER,
                               week_number
                               INTEGER,
                               tasks_completed
                               INTEGER,
                               total_tasks
                               INTEGER,
                               progress_score
                               REAL,
                               date
                               TEXT
                               DEFAULT
                               CURRENT_TIMESTAMP
                           );
                           CREATE INDEX IF NOT EXISTS idx_materials_sec_cat_book ON materials(section, category, book);
                           CREATE INDEX IF NOT EXISTS idx_leaderboard_weekly ON leaderboard(weekly_score DESC);
                           CREATE INDEX IF NOT EXISTS idx_daily_tasks_tid ON daily_tasks(telegram_id);
                           CREATE INDEX IF NOT EXISTS idx_history_tid ON history(telegram_id);
                           """)


def add_user(tid, username, full_name):
    with db() as conn:
        conn.execute(
            "INSERT OR IGNORE INTO users (telegram_id, username, full_name) VALUES (?,?,?)",
            (tid, username, full_name)
        )


def is_premium(tid):
    with db() as conn:
        row = conn.execute("SELECT is_premium FROM users WHERE telegram_id=?", (tid,)).fetchone()
        return bool(row and row[0])


def set_premium(tid, status):
    with db() as conn:
        conn.execute("UPDATE users SET is_premium=? WHERE telegram_id=?", (status, tid))


def get_user(tid):
    with db() as conn:
        return conn.execute(
            "SELECT telegram_id, username, full_name, is_premium FROM users WHERE telegram_id=?", (tid,)
        ).fetchone()


def get_all_users():
    with db() as conn:
        return conn.execute("SELECT telegram_id FROM users").fetchall()


def get_users_info():
    with db() as conn:
        return conn.execute(
            "SELECT telegram_id, username, full_name, is_premium, joined_date "
            "FROM users ORDER BY joined_date DESC LIMIT 20"
        ).fetchall()


def update_progress(tid, correct=False):
    with db() as conn:
        conn.execute("INSERT OR IGNORE INTO progress (telegram_id) VALUES (?)", (tid,))
        conn.execute(
            "UPDATE progress SET total_questions=total_questions+1, "
            "correct_answers=correct_answers+? WHERE telegram_id=?",
            (1 if correct else 0, tid)
        )
    if correct:
        add_leaderboard_score(tid, 10)


def get_progress(tid):
    with db() as conn:
        return conn.execute(
            "SELECT total_questions, correct_answers FROM progress WHERE telegram_id=?", (tid,)
        ).fetchone()


def save_material(section, category, book, test, part, title, file_id, file_type="document"):
    with db() as conn:
        conn.execute(
            "INSERT INTO materials (section,category,book,test,part,title,file_id,file_type) "
            "VALUES (?,?,?,?,?,?,?,?)",
            (section, category, book, test, part, title, file_id, file_type)
        )


def get_books(section, category):
    with db() as conn:
        rows = conn.execute(
            "SELECT DISTINCT book FROM materials WHERE section=? AND category=? ORDER BY book",
            (section, category)
        ).fetchall()
        return [r[0] for r in rows]


def get_tests(section, category, book):
    with db() as conn:
        rows = conn.execute(
            "SELECT DISTINCT test FROM materials WHERE section=? AND category=? AND book=? ORDER BY test",
            (section, category, book)
        ).fetchall()
        return [r[0] for r in rows]


def get_parts(section, category, book, test):
    with db() as conn:
        rows = conn.execute(
            "SELECT DISTINCT part FROM materials WHERE section=? AND category=? AND book=? AND test=? ORDER BY part",
            (section, category, book, test)
        ).fetchall()
        return [r[0] for r in rows]


def get_file(section, category, book, test, part):
    with db() as conn:
        return conn.execute(
            "SELECT id, title, file_id, file_type FROM materials "
            "WHERE section=? AND category=? AND book=? AND test=? AND part=?",
            (section, category, book, test, part)
        ).fetchone()


def get_all_materials():
    with db() as conn:
        return conn.execute(
            "SELECT id, section, category, book, test, part, title FROM materials ORDER BY date DESC LIMIT 50"
        ).fetchall()


def get_movies():
    with db() as conn:
        return conn.execute(
            "SELECT id, title, file_id, file_type FROM materials "
            "WHERE section='Movies' ORDER BY date DESC LIMIT 50"
        ).fetchall()


def get_books_list():
    with db() as conn:
        return conn.execute(
            "SELECT id, title, file_id, file_type FROM materials "
            "WHERE section='Books' ORDER BY date DESC"
        ).fetchall()


def delete_material(mid):
    with db() as conn:
        conn.execute("DELETE FROM materials WHERE id=?", (mid,))


def get_stats():
    with db() as conn:
        total = conn.execute("SELECT COUNT(*) FROM users").fetchone()[0]
        premium = conn.execute("SELECT COUNT(*) FROM users WHERE is_premium=1").fetchone()[0]
        mats = conn.execute("SELECT COUNT(*) FROM materials").fetchone()[0]
        qs = conn.execute("SELECT SUM(total_questions) FROM progress").fetchone()[0] or 0
        return total, premium, mats, qs


def save_history(tid, section, question, answer):
    with db() as conn:
        conn.execute(
            "INSERT INTO history (telegram_id,section,question,answer) VALUES (?,?,?,?)",
            (tid, section, question, answer)
        )


# ================== LEADERBOARD ==================

def add_leaderboard_score(tid, points):
    with db() as conn:
        conn.execute("INSERT OR IGNORE INTO leaderboard (telegram_id) VALUES (?)", (tid,))
        conn.execute(
            "UPDATE leaderboard SET weekly_score=weekly_score+?, total_score=total_score+?, "
            "last_updated=CURRENT_TIMESTAMP WHERE telegram_id=?",
            (points, points, tid)
        )


def get_leaderboard(limit=10):
    with db() as conn:
        return conn.execute(
            "SELECT l.telegram_id, u.full_name, u.username, l.weekly_score, l.total_score "
            "FROM leaderboard l JOIN users u ON l.telegram_id=u.telegram_id "
            "ORDER BY l.weekly_score DESC LIMIT ?",
            (limit,)
        ).fetchall()


def get_my_rank(tid):
    with db() as conn:
        score_row = conn.execute("SELECT weekly_score FROM leaderboard WHERE telegram_id=?", (tid,)).fetchone()
        if not score_row:
            return None
        rank = conn.execute("SELECT COUNT(*) FROM leaderboard WHERE weekly_score > ?", (score_row[0],)).fetchone()[0]
        return rank + 1


def reset_weekly_scores():
    with db() as conn:
        conn.execute("UPDATE leaderboard SET weekly_score=0")
    print("✅ Haftalik ballar nollandi!")


# ================== STUDY PLAN FUNCTIONS ==================

def save_study_plan(tid, current_level, target_level, duration_weeks, plan_content):
    with db() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO study_plans (telegram_id, current_level, target_level, duration_weeks, plan_content) VALUES (?,?,?,?,?)",
            (tid, current_level, target_level, duration_weeks, plan_content)
        )

def get_study_plan(tid):
    with db() as conn:
        return conn.execute(
            "SELECT current_level, target_level, duration_weeks, plan_content, created_date FROM study_plans WHERE telegram_id=? ORDER BY created_date DESC LIMIT 1",
            (tid,)
        ).fetchone()

def update_weekly_progress(tid, week_number, tasks_completed, total_tasks, progress_score):
    with db() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO weekly_progress (telegram_id, week_number, tasks_completed, total_tasks, progress_score) VALUES (?,?,?,?,?)",
            (tid, week_number, tasks_completed, total_tasks, progress_score)
        )

def get_weekly_progress(tid):
    with db() as conn:
        return conn.execute(
            "SELECT week_number, tasks_completed, total_tasks, progress_score FROM weekly_progress WHERE telegram_id=? ORDER BY week_number",
            (tid,)
        ).fetchall()


def create_study_plan_pdf(current_level, target_level, duration_weeks, weak_skills,
                          plan_content, design_choice="minimal", study_hours="2 soat"):
    """IELTS Study Plan PDF — Professional Design (barcha xatolar tuzatilgan)"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        from reportlab.lib.colors import HexColor
        import io

        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4,
                                topMargin=0.5 * inch, bottomMargin=0.5 * inch,
                                leftMargin=0.6 * inch, rightMargin=0.6 * inch)
        styles = getSampleStyleSheet()

        # Ranglar
        if design_choice == "creative":
            primary   = HexColor('#6C3483')   # binafsha
            secondary = HexColor('#F9EBFF')
            accent    = HexColor('#1ABC9C')
        else:  # minimal
            primary   = HexColor('#2E86AB')   # ko'k
            secondary = HexColor('#E8F4FD')
            accent    = HexColor('#28A745')

        white = colors.white
        black = colors.black

        title_style = ParagraphStyle('T', parent=styles['Title'],
                                     fontSize=24, alignment=TA_CENTER,
                                     textColor=white, fontName='Helvetica-Bold')
        section_style = ParagraphStyle('S', parent=styles['Heading2'],
                                       fontSize=13, textColor=primary,
                                       fontName='Helvetica-Bold', spaceAfter=6)
        body_style = ParagraphStyle('B', parent=styles['Normal'],
                                    fontSize=10, leading=14,
                                    fontName='Helvetica', spaceAfter=4)

        story = []

        # ── SARLAVHA ──
        header_data = [[Paragraph("🏆 ELITE PREP — IELTS STUDY PLAN", title_style)]]
        header_table = Table(header_data, colWidths=[7 * inch])
        header_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), primary),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('PADDING', (0, 0), (-1, -1), 18),
        ]))
        story.append(header_table)
        story.append(Spacer(1, 16))

        # ── PROFIL ──
        story.append(Paragraph("👤 STUDENT PROFILE", section_style))
        weak_str = ', '.join(weak_skills).title() if weak_skills else 'All Skills (Balanced)'
        profile_data = [
            ["Current Band", "Target Band", "Duration", "Daily Study", "Focus"],
            [f"{current_level}/9.0", f"{target_level}/9.0",
             f"{duration_weeks} weeks", study_hours, weak_str],
        ]
        col_w = [1.3 * inch, 1.2 * inch, 1.2 * inch, 1.2 * inch, 2.1 * inch]
        pt = Table(profile_data, colWidths=col_w)
        pt.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), primary),
            ('BACKGROUND', (0, 1), (-1, 1), secondary),
            ('TEXTCOLOR', (0, 0), (-1, 0), white),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('FONTSIZE', (0, 1), (-1, 1), 11),
            ('FONTNAME', (0, 1), (-1, 1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, primary),
            ('PADDING', (0, 0), (-1, -1), 9),
        ]))
        story.append(pt)
        story.append(Spacer(1, 18))

        # ── AI PLAN MATNI ──
        story.append(Paragraph("📅 PERSONALIZED WEEKLY PLAN", section_style))

        # Plan matnini chiroyli formatlash
        for line in plan_content.split('\n'):
            stripped = line.strip()
            if not stripped:
                story.append(Spacer(1, 4))
                continue

            # Hafta sarlavhasi
            if stripped.startswith('━') or ('WEEK' in stripped.upper() and '━' in stripped):
                story.append(Spacer(1, 8))
                week_data = [[Paragraph(f"<b>{stripped}</b>",
                                        ParagraphStyle('WH', parent=styles['Normal'],
                                                       textColor=white, fontSize=11,
                                                       fontName='Helvetica-Bold'))]]
                wt = Table(week_data, colWidths=[7 * inch])
                wt.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, -1), primary),
                    ('PADDING', (0, 0), (-1, -1), 8),
                ]))
                story.append(wt)
                story.append(Spacer(1, 4))

            # Ko'nikmalar
            elif any(stripped.startswith(k) for k in ['📖', '✍️', '🎧', '🗣', '▸', '•', '🎯', '📈', '📊']):
                story.append(Paragraph(stripped, ParagraphStyle('KS', parent=styles['Normal'],
                                                                  fontSize=10, leading=14,
                                                                  fontName='Helvetica',
                                                                  leftIndent=10)))
            else:
                story.append(Paragraph(stripped, body_style))

        story.append(Spacer(1, 20))

        # ── FOOTER ──
        footer_data = [[Paragraph(
            "© ELITE PREP AI | Your Personal IELTS Mentor | Telegram: @eliteprep_bot",
            ParagraphStyle('F', parent=styles['Normal'],
                           alignment=TA_CENTER, fontSize=8,
                           textColor=white, fontName='Helvetica-Oblique')
        )]]
        ft = Table(footer_data, colWidths=[7 * inch])
        ft.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), primary),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(ft)

        doc.build(story)
        buffer.seek(0)
        return buffer

    except Exception as e:
        print(f"PDF xatosi: {e}")
        import traceback
        traceback.print_exc()
        return None


# ─────────────────────────────────────────────
# 2. YORDAMCHI — IELTS ball formati
# ─────────────────────────────────────────────

def format_ielts_score(score: float) -> float:
    """
    IELTS ballini faqat .0 yoki .5 formatiga yaxlitlaydi.
    Masalan: 6.3 → 6.5,  6.8 → 7.0
    """
    whole = int(score)
    decimal = score - whole
    if decimal < 0.25:
        return float(whole)
    elif decimal < 0.75:
        return float(whole + 0.5)
    else:
        return float(whole + 1.0)


# ─────────────────────────────────────────────
# 3. AI STUDY PLAN GENERATOR — takomillashtirilgan
# ─────────────────────────────────────────────

def ai_generate_study_plan(current_level, target_level, duration_weeks, study_hours, weak_skills=None):
    """
    AI yordamida shaxsiy IELTS o'quv rejasi yaratadi.
    Weak skills foydalanuvchidan olingan bo'ladi.
    """
    if weak_skills is None:
        weak_skills = []

    # Har bir ko'nikmaga vaqt ulushi (%)
    base_pct = 25
    skill_focus = {"reading": base_pct, "writing": base_pct,
                   "listening": base_pct, "speaking": base_pct}

    if weak_skills:
        boost = 15
        for skill in weak_skills:
            if skill in skill_focus:
                skill_focus[skill] += boost
        # Qolganlarni kamaytiramiz
        total_extra = boost * len(weak_skills)
        non_weak = [s for s in skill_focus if s not in weak_skills]
        if non_weak:
            cut = total_extra // len(non_weak)
            for s in non_weak:
                skill_focus[s] = max(10, skill_focus[s] - cut)

    prompt = f"""You are an elite IELTS tutor with 15+ years of experience.

STUDENT PROFILE:
• Current Band: {current_level}/9.0
• Target Band: {target_level}/9.0
• Duration: {duration_weeks} weeks
• Daily Study Time: {study_hours}
• Weak Areas: {', '.join(weak_skills).title() if weak_skills else 'Balanced across all skills'}

TIME ALLOCATION:
📖 Reading: {skill_focus['reading']}%
✍️ Writing: {skill_focus['writing']}%
🎧 Listening: {skill_focus['listening']}%
🗣 Speaking: {skill_focus['speaking']}%

Create a DETAILED, WEEKLY study plan in the following format for ALL {duration_weeks} weeks:

━━━━━━━━━━━━━━━━━━━━
📅 WEEK [N] — [Focus Theme]
━━━━━━━━━━━━━━━━━━━━
🎯 Weekly Target: [specific measurable goal]
📈 Expected Band After This Week: [X]/9.0

DAILY SCHEDULE ({study_hours}/day):
▸ Monday:    [Task — Resource — Duration]
▸ Tuesday:   [Task — Resource — Duration]
▸ Wednesday: [Task — Resource — Duration]
▸ Thursday:  [Task — Resource — Duration]
▸ Friday:    [Task — Resource — Duration]
▸ Saturday:  FULL MOCK TEST + Review
▸ Sunday:    Error analysis + Rest

SKILL FOCUS:
📖 Reading:   [Specific strategy + Cambridge test number]
✍️ Writing:   [Task 1 or Task 2 topic + word count]
🎧 Listening: [Section type + Cambridge test number]
🗣 Speaking:  [Part 1/2/3 topic + cue card]

📊 Progress Check:
• Target score this week: [X]/9.0
• Key improvement area: [specific tip]

(Repeat the above block for each of the {duration_weeks} weeks)

━━━━━━━━━━━━━━━━━━━━
🏆 FINAL ADVICE & SUCCESS TIPS
━━━━━━━━━━━━━━━━━━━━
[3-4 motivational and practical tips]

Write in ENGLISH. Be specific, realistic, and motivational!"""

    try:
        if groq:
            resp = groq.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[
                    {"role": "system",
                     "content": "You are the world's best IELTS coach. Create highly specific, week-by-week study plans."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=4000
            )
            return resp.choices[0].message.content
        elif gemini:
            resp = gemini.models.generate_content(model="gemini-2.0-flash", contents=prompt)
            return resp.text
    except Exception as e:
        print(f"AI Study Plan xato: {e}")

    # AI ishlamasa fallback
    return _fallback_study_plan(current_level, target_level, duration_weeks, study_hours, weak_skills)


def _fallback_study_plan(current_level, target_level, duration_weeks, study_hours, weak_skills):
    """AI ishlamasa ishlaydigan zaxira reja."""
    weak_str = ', '.join(weak_skills).title() if weak_skills else 'Barcha ko\'nikmalar'
    lines = [
        f"📊 IELTS STUDY PLAN: {current_level} → {target_level}",
        f"⏰ Kunlik: {study_hours} | 📅 Muddat: {duration_weeks} hafta",
        f"🎯 Zaif tomonlar: {weak_str}",
        ""
    ]
    for w in range(1, duration_weeks + 1):
        lines += [
            f"━━━ HAFTA {w} ━━━",
            "Dushanba:   Listening — Cambridge Test",
            "Seshanba:   Reading — Academic Passage",
            "Chorshanba: Writing — Task 2 Essay (250+ so'z)",
            "Payshanba:  Speaking — Part 1 & 2",
            "Juma:       Vocabulary + Grammar",
            "Shanba:     To'liq Mock Test",
            "Yakshanba:  Xatolarni tahlil qilish + Hafta yakuni"
        ]
    
    return "\n".join(lines)

# ================== STATE HELPERS ==================

def get_state(tid):
    with state_lock:
        return user_states.get(tid)


def set_state(tid, state):
    with state_lock:
        user_states[tid] = state


def del_state(tid):
    with state_lock:
        user_states.pop(tid, None)


# ================== FAYL YUBORISH ==================

def send_file(chat_id, file_id, file_type, caption=""):
    try:
        if file_type == "audio":
            bot.send_audio(chat_id, file_id, caption=caption)
        elif file_type == "video":
            bot.send_video(chat_id, file_id, caption=caption)
        else:
            bot.send_document(chat_id, file_id, caption=caption)
    except Exception as e:
        bot.send_message(chat_id, f"❌ Fayl yuborishda xato: {e}")


# ================== AI FUNKSIYALAR ==================

def ai_speaking(user_text):
    resp = groq.chat.completions.create(
        model="llama-3.3-70b-versatile",
        messages=[
            {"role": "system", "content": (
                "You are Sarah, a friendly IELTS Speaking coach. "
                "Reply in 2-3 natural sentences + 1 follow-up question. "
                "Be warm and encouraging. "
                "If grammar mistakes, add: [Tuzatish: explanation in Uzbek]. "
                "Never say 'As an AI'. Use: Oh!, Really?, Wow!"
            )},
            {"role": "user", "content": user_text}
        ],
        temperature=0.9,
        max_tokens=200
    )
    return resp.choices[0].message.content


def ai_writing(essay):
    prompt = "You are a world-class IELTS Writing Examiner. Analyze this essay:\n\n" + \
             f"ESSAY: {essay}\n\n" + \
             "Provide feedback in English with these sections:\n" + \
             "1. OVERALL BAND SCORE (4-9 realistic range)\n" + \
             "2. CRITICAL STRENGTHS\n" + \
             "3. KEY WEAKNESSES\n" + \
             "4. DETAILED CRITERIA:\n" + \
             "   - Task Response (TR)\n" + \
             "   - Coherence and Cohesion (CC)\n" + \
             "   - Lexical Resource (LR)\n" + \
             "   - Grammatical Range and Accuracy (GRA)\n" + \
             "5. STRUCTURAL ADVICE\n" + \
             "6. IMPROVED VERSION (clean text, no formatting)\n\n" + \
             "IMPORTANT: Provide realistic IELTS scores (4.0-9.0). Never give scores below 4.0 unless the essay is completely incomprehensible.\n\n" + \
             "Focus on IELTS standards."

    try:
        if groq:
            resp = groq.chat.completions.create(
                model="llama-3.3-70b-versatile",
                messages=[{"role": "system", "content": "You are a senior IELTS Writing Examiner."},
                          {"role": "user", "content": prompt}],
                temperature=0.7, max_tokens=4000
            )
            return resp.choices[0].message.content
        elif gemini:
            resp = gemini.models.generate_content(model="gemini-2.0-flash", contents=prompt)
            return resp.text
    except Exception as e:
        print(f"AI Writing xato: {e}")
    return "⚠️ AI service is currently unavailable."


def create_writing_report_docx(essay_text, ai_analysis):
    # Professional IELTS Writing Report Generator.
    # - Improved Version with numbered topic vocabulary list (with Uzbek translation)
    # - Band Score after Candidate Submission
    # - Only 4 detailed criteria (TR, CC, LR, GRA) in bold
    doc = Document()

    # ── Global Style: Times New Roman 14pt ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(14)

    # ── HEADER: Title ──
    title = doc.add_heading('IELTS PREP AI — ASSESSMENT REPORT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFB, 0x71, 0x85)

    # ── Meta info ──
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r1 = p_meta.add_run(f'Report ID: #WD{int(time.time())}\n')
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r2 = p_meta.add_run(f'Date: {datetime.now().strftime("%B %d, %Y")}\n')
    r2.font.name = 'Times New Roman'
    r3 = p_meta.add_run('ElitePrep AI — Professional Writing Evaluation')
    r3.italic = True
    r3.font.name = 'Times New Roman'

    doc.add_paragraph()

    # ── CANDIDATE ESSAY ──
    h_essay = doc.add_heading('Candidate Submission:', level=1)
    for run in h_essay.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    orig_p = doc.add_paragraph(essay_text)
    orig_p.paragraph_format.space_after = Pt(6)
    for run in orig_p.runs:
        run.font.name = 'Times New Roman'
    
    # Overall band score ni Candidate Submission dan keyin qo'yish
    import re
    band_score = ""
    for line in ai_analysis.split('\n'):
        upper = line.upper()
        if 'OVERALL BAND SCORE' in upper:
            cleaned = line.strip().replace('#', '').replace('*', '').strip()
            # 4-9 oralig'idagi ballni qidirish (masalan: 7.5, 8, 6.0)
            m = re.search(r'\b[4-9](?:\.\d+)?\b', cleaned)
            if m:
                score = m.group()
                # Faqat 4-9 oralig'idagi ballarni qabul qilish
                if 4 <= float(score) <= 9:
                    band_score = score
            break # Ball topilgach, qidirishni to'xtatish
    
    if band_score:
        doc.add_paragraph()
        p_band = doc.add_paragraph()
        p_band.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_band = p_band.add_run(f'Overall band score {band_score}')
        run_band.bold = True
        run_band.font.size = Pt(16)
        run_band.font.name = 'Times New Roman'
        run_band.font.color.rgb = RGBColor(0xFB, 0x71, 0x85)
        p_band.paragraph_format.space_before = Pt(12)
        p_band.paragraph_format.space_after = Pt(12)

    doc.add_paragraph()

    # ── EXAMINER ANALYSIS ──
    h_analysis = doc.add_heading('Examiner Analysis & Feedback:', level=1)
    for run in h_analysis.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

    # Faqat 4 ta criteria — qalin yozuv
    CRITERIA_4_ONLY = ['TASK RESPONSE', 'COHERENCE AND COHESION', 'LEXICAL RESOURCE',
                       'GRAMMATICAL RANGE AND ACCURACY']

    SECTION_KEYWORDS = ['OVERALL BAND SCORE', 'CRITICAL STRENGTHS', 'KEY WEAKNESSES',
                        'DETAILED CRITERIA', 'STRUCTURAL ADVICE', 'IMPROVED VERSION',
                        'KEY VOCABULARY', 'SUGGESTIONS', 'EVALUATION']

    inside_detailed_criteria = False
    skip_to_next_section = False

    for line in ai_analysis.split('\n'):
        raw = line.strip()
        if not raw:
            doc.add_paragraph()
            continue

        clean = raw.replace('#', '').replace('**', '').strip()
        upper = clean.upper()

        # Section check
        is_section = any(kw in upper for kw in SECTION_KEYWORDS)
        is_criteria = any(kw in upper for kw in CRITERIA_4_ONLY)

        # Track if we're in detailed criteria
        if 'DETAILED CRITERIA' in upper:
            inside_detailed_criteria = True
        elif is_section and 'DETAILED CRITERIA' not in upper:
            inside_detailed_criteria = False

        # Skip "OVERALL BAND SCORE" section — allaqachon chiqarilgan
        if 'OVERALL BAND SCORE' in upper:
            skip_to_next_section = True
            continue

        # Skip "IMPROVED VERSION" section — oxiriga qo'yiladi
        if 'IMPROVED VERSION' in upper:
            skip_to_next_section = True
            continue
        if skip_to_next_section and is_section:
            skip_to_next_section = False

        if skip_to_next_section:
            continue

        if is_section:
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(4)
            
            # Section raqamlarini to'g'rilash
            if 'CRITICAL STRENGTHS' in upper:
                clean = '1. CRITICAL STRENGTHS'
            elif 'KEY WEAKNESSES' in upper:
                clean = '2. KEY WEAKNESSES'
            elif 'DETAILED CRITERIA' in upper:
                clean = '3. DETAILED CRITERIA'
            elif 'STRUCTURAL ADVICE' in upper:
                clean = '4. STRUCTURAL ADVICE'
            else:
                clean = clean  # Boshqa sectionlarni o'zgartirmaslik
            
            run = para.add_run(clean)
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        elif is_criteria and inside_detailed_criteria:
            # 4 ta criteria — qalin yozuv
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(4)
            run = para.add_run(clean)
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        else:
            # Oddiy matn
            para = doc.add_paragraph(clean)
            para.paragraph_format.space_after = Pt(4)
            for r in para.runs:
                r.font.name = 'Times New Roman'

    # ── IMPROVED VERSION ──
    doc.add_paragraph()

    # Improved version'ni topish
    improved_text = ""
    capture_improved = False
    skip_first_line = True
    for line in ai_analysis.split('\n'):
        upper = line.upper()
        if 'IMPROVED VERSION' in upper and 'KEY VOCABULARY' not in upper:
            capture_improved = True
            skip_first_line = True
            continue
        if capture_improved:
            if 'KEY VOCABULARY' in upper:
                break
            # Birinchi bo'sh yoki heading lineni skip qil
            if skip_first_line and (not line.strip() or 'IMPROVED' in line.upper()):
                skip_first_line = False
                continue
            skip_first_line = False
            improved_text += line + "\n"

    if improved_text.strip():
        # Improved Version heading qo'shish
        h_improved = doc.add_heading('Improved Version:', level=1)
        for run in h_improved.runs:
            run.font.name = 'Times New Roman'
            run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        
        para_imp = doc.add_paragraph(improved_text.strip())
        para_imp.paragraph_format.space_after = Pt(10)
        for r in para_imp.runs:
            r.font.name = 'Times New Roman'

    # ── FOOTER ──
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.text = "© IELTS Prep AI | Your Personal IELTS Mentor | Telegram: @ielts_prep_ai_bot"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(10)
        run.font.italic = True
        run.font.name = 'Times New Roman'

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream


def send_voice(chat_id, text):
    english = text.split("[Tuzatish:")[0].strip()
    try:
        if eleven:
            audio = eleven.text_to_speech.convert(
                voice_id="21m00Tcm4TlvDq8ikWAM",
                text=english,
                model_id="eleven_turbo_v2"
            )
            with open(f"resp_{chat_id}.mp3", "wb") as f:
                for chunk in audio:
                    f.write(chunk)
            with open(f"resp_{chat_id}.mp3", "rb") as f:
                bot.send_voice(chat_id, f)
        else:
            raise RuntimeError("ElevenLabs kaliti yo'q")
    except Exception:
        try:
            tts = gTTS(text=english, lang="en")
            tts.save(f"resp_{chat_id}.mp3")
            with open(f"resp_{chat_id}.mp3", "rb") as f:
                bot.send_voice(chat_id, f)
        except Exception:
            bot.send_message(chat_id, f"🔊 {english}")
    finally:
        try:
            os.remove(f"resp_{chat_id}.mp3")
        except:
            pass


# ================== MOCK TEST MA'LUMOTLARI ==================
# BOTDA FULL MOCK TEST ISHLATISH UCHUN SAVOLLARNI O'ZINGIZ QO'SHISHINGIZ MUMKIN
MOCK_TEST_DATA = {
    # LISTENING UCHUN SAVOLLAR (Bu yerga audioni telegramda yuborib joyini tushuntirsangiz bo'ladi)
    "listening": [
        {
            "q": "Listening 1-savol: The man wants to buy a...", 
            "options": {"A": "Car", "B": "Bike", "C": "House", "D": "Boat"}, 
            "correct": "A"
        },
        {
            "q": "Listening 2-savol: The color of the bike is...", 
            "options": {"A": "Red", "B": "Blue", "C": "Green", "D": "Black"}, 
            "correct": "B"
        },
        # Shu tartibda savol qo'shasiz (namuna sifatida 2 ta e'lon qilingan)
    ],
    # READING UCHUN SAVOLLAR (Bu yerga Reading Passage va savollarni qo'shing)
    "reading": [
        {
            "q": "Reading passage 1\nThere are many advantages to reading everyday. According to the text...", 
            "options": {"A": "Yes", "B": "No", "C": "Not Given", "D": "None"}, 
            "correct": "C"
        },
        {
            "q": "Reading 2-savol: The author states that...", 
            "options": {"A": "Option A", "B": "Option B", "C": "Option C", "D": "Option D"}, 
            "correct": "A"
        },
        # Shu tartibda savol qo'shasiz
    ],
    # WRITING UCHUN MAVZU
    "writing": "Some people think technology has made our lives more complicated. Do you agree or disagree?\n(Kamida 250 so'z)",
    # SPEAKING UCHUN SAVOLLAR
    "speaking": [
        "Tell me about your hometown. What do you like about it?",
        "What do you enjoy doing in your free time?",
        "Describe your favourite food and why you like it?",
        "What are your plans for the future?",
        "How important is learning English in your life?"
    ]
}


def mock_question(chat_id, tid, section, q_num):
    if section == "writing":
        bot.send_message(chat_id,
                         f"✍️ *Writing — Essay yozing*\n\n"
                         f"*Mavzu:* _{MOCK_TEST_DATA['writing']}_\n\n_(Kamida 250 so'z)_",
                         parse_mode="Markdown")
        return
    if section == "speaking":
        questions = MOCK_TEST_DATA.get("speaking", [])
        if q_num > len(questions):
            return
        bot.send_message(chat_id,
                         f"🗣 *Speaking — Savol {q_num}/{len(questions)}*\n\n_{questions[q_num - 1]}_\n\nJavob bering!",
                         parse_mode="Markdown")
        return

    questions = MOCK_TEST_DATA.get(section, [])
    if q_num > len(questions) or len(questions) == 0:
        bot.send_message(chat_id, "⚠️ Hozircha bu bo'limda savollar yo'q.")
        return

    q_data = questions[q_num - 1]
    text = q_data["q"]
    correct = q_data["correct"]
    
    markup = telebot.types.InlineKeyboardMarkup()
    buttons = []
    for opt_key, opt_val in q_data["options"].items():
        text += f"\n{opt_key}) {opt_val}"
        buttons.append(telebot.types.InlineKeyboardButton(opt_key, callback_data=f"mock_{opt_key}_{correct}_{q_num}_{section}"))
    
    markup.row(*buttons)
    icons = {"listening": "🎧", "reading": "📖"}
    bot.send_message(chat_id,
                     f"{icons.get(section, '')} *Savol {q_num}/{len(questions)}*\n\n{text}",
                     parse_mode="Markdown", reply_markup=markup)


def finish_mock(tid):
    state = get_state(tid) or {}
    answers = state.get("answers", {})
    l = answers.get("listening", {})
    r = answers.get("reading", {})
    ls = round((l.get("correct", 0) / max(l.get("total", 1), 1)) * 9, 1)
    rs = round((r.get("correct", 0) / max(r.get("total", 1), 1)) * 9, 1)
    overall = round((ls + rs) / 2, 1)
    add_leaderboard_score(tid, int(overall * 10))
    bot.send_message(tid,
                     f"🎉 *Full Mock Test yakunlandi!*\n\n"
                     f"🎧 Listening: *{ls}/9*\n"
                     f"📖 Reading: *{rs}/9*\n"
                     f"✍️ Writing: *✅*\n"
                     f"🗣 Speaking: *✅*\n\n"
                     f"📊 *Taxminiy ball: {overall}/9*\n\n"
                     f"🏆 Leaderboard balling yangilandi!\n"
                     f"💪 Davom eting!",
                     parse_mode="Markdown", reply_markup=main_menu())
    del_state(tid)


# ================== MENYULAR ==================

def main_menu():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row("🎓 IELTS", "🏫 SAT")
    m.row("📚 Grammar & Vocab")
    m.row("🎬 Movies", "📗 Books")
    m.row(" Premium")
    return m


def ai_consultant_menu():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    # placement test removed per user request; study plan also removed
    m.row("🔙 Orqaga")
    return m


def ielts_menu():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row("📖 Reading", "🎧 Listening")
    m.row("✍️ Writing", "🗣 Speaking")
    # web app ishlatish!
    if MOCK_WEBAPP_URL:
        m.row(telebot.types.KeyboardButton("💻 Full Mock Test (Web)", web_app=telebot.types.WebAppInfo(url=MOCK_WEBAPP_URL)))
    else:
        m.row("📝 Full Mock Test")
    m.row("🔙 Orqaga")
    return m


def category_menu(section):
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row(f"📚 Cambridge|{section}")
    m.row(f"⭐ Ultimate|{section}")
    m.row("🔙 Orqaga")
    return m


def writing_menu():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row("🤖 AI Writing Checker")
    m.row("🔙 Orqaga")
    return m


def sat_menu():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row("📖 Reading & Writing", "📐 Math")
    m.row("📊 Ballim", "🔙 Orqaga")
    return m


def grammar_menu():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row("📝 Grammar", "📖 Vocabulary")
    m.row("🔙 Orqaga")
    return m


def speaking_menu():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    if WEBAPP_URL:
        m.row(telebot.types.KeyboardButton(
            "🤖 AI bilan mashq",
            web_app=telebot.types.WebAppInfo(url=WEBAPP_URL)
        ))
        m.row("👥 Odam bilan mashq")
    else:
        m.row("🤖 AI bilan mashq", "👥 Odam bilan mashq")
    m.row("🔙 Orqaga")
    return m


def admin_menu():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row("📊 Statistika", "👥 Foydalanuvchilar")
    m.row("💎 Premium berish", "📤 Xabar yuborish")
    m.row("📂 Materiallar", "🔙 Orqaga")
    return m


def cancel_markup():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row("❌ Bekor qilish")
    return m


def end_chat_markup():
    m = telebot.types.ReplyKeyboardMarkup(resize_keyboard=True)
    m.row("❌ Suhbatni tugatish")
    return m


# ================== INLINE HELPERS ==================

def tests_inline(section, category, book):
    tests = get_tests(section, category, book)
    m = telebot.types.InlineKeyboardMarkup()
    row = []
    for t in tests:
        row.append(telebot.types.InlineKeyboardButton(t, callback_data=f"test_{section}_{category}_{book}_{t}"))
        if len(row) == 2:
            m.row(*row);
            row = []
    if row: m.row(*row)
    return m, tests


def parts_inline(section, category, book, test):
    parts = get_parts(section, category, book, test)
    m = telebot.types.InlineKeyboardMarkup()
    row = []
    for p in parts:
        row.append(telebot.types.InlineKeyboardButton(p, callback_data=f"part_{section}_{category}_{book}_{test}_{p}"))
        if len(row) == 2:
            m.row(*row);
            row = []
    if row: m.row(*row)
    return m, parts


# ================== HANDLERS ==================

@bot.message_handler(commands=["start"])
def cmd_start(msg):
    add_user(msg.from_user.id, msg.from_user.username, msg.from_user.full_name)
    set_reminder(msg.from_user.id, is_active=True)
    name = msg.from_user.first_name or "Do'st"
    bot.send_message(msg.chat.id,
                     f"👋 Salom, *{name}*!\n\n"
                     f"📚 *ElitePrep AI* ga xush kelibsiz!\n\n"
                     f"✅ IELTS va SAT tayyorlanish\n"
                     f"✅ Cambridge materiallar\n"
                     f"✅ AI Writing Checker\n"
                     f"✅ AI Speaking mashqi\n"
                     f"✅ Full Mock Test\n"
                     f"✅ Kunlik vazifalar 📅\n"
                     f"✅ Leaderboard 🏆\n\n"
                     f"Boshlaylik! 👇",
                     parse_mode="Markdown", reply_markup=main_menu())


@bot.message_handler(commands=["admin"])
def cmd_admin(msg):
    if msg.from_user.id != ADMIN_ID:
        bot.send_message(msg.chat.id, "❌ Ruxsat yo'q!")
        return
    bot.send_message(msg.chat.id, "👑 *Admin Panel*", parse_mode="Markdown", reply_markup=admin_menu())


@bot.message_handler(commands=["stop"])
def cmd_stop(msg):
    del_state(msg.from_user.id)
    bot.send_message(msg.chat.id, "✅ Tugadi!", reply_markup=main_menu())


# ================== FAYL HANDLER ==================

@bot.message_handler(content_types=["document", "audio", "video", "photo"])
def handle_file(msg):
    tid = msg.from_user.id
    if tid != ADMIN_ID:
        return
    
    file_id = None
    file_type = ""
    
    if msg.content_type == "document":
        file_id = msg.document.file_id; file_type = "document"
    elif msg.content_type == "audio":
        file_id = msg.audio.file_id; file_type = "audio"
    elif msg.content_type == "video":
        file_id = msg.video.file_id; file_type = "video"
    elif msg.content_type == "photo":
        file_id = msg.photo[-1].file_id; file_type = "photo"

    # ✅ Media Mock handle
    state = get_state(tid)
    if isinstance(state, dict) and state.get("mode") == "media_mock":
        f_info = bot.get_file(file_id)
        ext = "mp3" if file_type == "audio" else "png"
        f_path = f"tmp_{tid}_{state['step']}_{len(state.get('files', []))}.{ext}"
        url = f"https://api.telegram.org/file/bot{BOT_TOKEN}/{f_info.file_path}"
        import urllib.request
        urllib.request.urlretrieve(url, f_path)
        
        if "files" not in state: state["files"] = []
        state["files"].append({"step": state["step"], "path": f_path})
        set_state(tid, state)
        bot.send_message(tid, f"✅ {state['step'].upper()} fayli qabul qilindi. Yana yuboring yoki '✅ Keyingisi' bosing.")
        return

    caption = msg.caption
    if not caption:
        if file_type == "photo": return
        bot.send_message(tid, "⚠️ *Format:*\n`Section|Category|Book|Test|Part|Title`", parse_mode="Markdown")
        return

    parts = [p.strip() for p in caption.split("|")]

    if len(parts) == 2 and parts[0].lower() == "movies":
        save_material("Movies", "Movies", "Movies", "Movies", "Movies", parts[1], file_id, file_type)
        bot.send_message(msg.chat.id, f"✅ *Movie saqlandi!*\n\n🎬 {parts[1]}", parse_mode="Markdown")

    elif len(parts) == 2 and parts[0].lower() == "books":
        save_material("Books", "Books", "Books", "Books", "Books", parts[1], file_id, file_type)
        bot.send_message(msg.chat.id, f"✅ *Kitob saqlandi!*\n\n📗 {parts[1]}", parse_mode="Markdown")

    # ✅ NEW: Media Mock handle
    state = get_state(msg.from_user.id)
    if isinstance(state, dict) and state.get("mode") == "media_mock":
        bot.send_message(msg.chat.id, "✅ Fayl qabul qilindi. Yana yuborishingiz mumkin yoki '✅ Keyingisi' deb yozing.")
        f_info = bot.get_file(file_id)
        f_path = f"tmp_{msg.from_user.id}_{state['step']}_{len(state.get('files', []))}.png"
        url = f"https://api.telegram.org/file/bot{BOT_TOKEN}/{f_info.file_path}"
        import urllib.request
        urllib.request.urlretrieve(url, f_path)
        
        if "files" not in state: state["files"] = []
        state["files"].append({"step": state["step"], "path": f_path})
        set_state(msg.from_user.id, state)
        return

    if len(parts) == 6:
        section, category, book, test, part, title = parts
        save_material(section, category, book, test, part, title, file_id, file_type)
        bot.send_message(msg.chat.id,
                         f"✅ *Saqlandi!*\n\n📂 {section} → {category} → {book}\n🧪 {test} → {part}\n🏷 {title}",
                         parse_mode="Markdown")
    else:
        bot.send_message(msg.chat.id,
                         "⚠️ Format noto'g'ri!\n\n"
                         "Materiallar: `Section|Category|Book|Test|Part|Title`\n"
                         "Filmlar: `Movies|Nomi`\n"
                         "Kitoblar: `Books|Nomi`",
                         parse_mode="Markdown")


# ================== MOCK WEB APP DATA HANDLER ==================

@bot.message_handler(content_types=['web_app_data'])
def handle_web_app_data(msg):
    import json
    try:
        data = json.loads(msg.web_app_data.data)
        if data.get("type") == "mock_test_results":
            tid = msg.from_user.id
            l_score = data.get("listening", {"correct": 0, "total": 1})
            r_score = data.get("reading", {"correct": 0, "total": 1})
            writing_text = data.get("writing", "")

            # Hisob-kitob qilish
            ls = round((l_score["correct"] / max(l_score["total"], 1)) * 9, 1)
            rs = round((r_score["correct"] / max(r_score["total"], 1)) * 9, 1)

            bot.send_message(tid, f"✅ *Reading & Listening hisoblandi!*\n🎧 L: *{ls}/9.0*  |  📖 R: *{rs}/9.0*\n\n⏳ Essay AI yordamida tekshirilmoqda...", parse_mode="Markdown")

            def process_results():
                try:
                    w_result = ai_writing(writing_text)
                    for i in range(0, len(w_result), 4000):
                        bot.send_message(tid, w_result[i:i+4000], parse_mode="Markdown")
                except Exception as e:
                    bot.send_message(tid, f"❌ Essay xatosi: {e}")
                
                # O'rgatish speaking speaking holatiga o'tkazish
                state = {"mode": "mock_speaking_ai", "speak_q": 1, "ls": ls, "rs": rs}
                set_state(tid, state)
                bot.send_message(tid, "✅ *Writing tugadi!*\n\n🗣 *Dinamik AI Speaking boshlanyapti...*\nAI Examiner ovozli xabarini kuting va mikrofonga yozib javob bering!", parse_mode="Markdown")

                # AI examiner boshlaydi
                ai_sys_prompt = (
                    "You are Sarah, a professional and strict IELTS Speaking Examiner. "
                    "Your mission is to conduct a complete 3-part speaking mock test. "
                    "PART 1: General questions. PART 2: Cue Card (tell user to speak for 2 minutes). PART 3: Abstract discussion based on Part 2. "
                    "Follow this flow: Ask 3 questions for Part 1, then give a Cue Card topic for Part 2, then ask 2-3 deep questions for Part 3. "
                    "Only ask ONE question at a time. After Part 3 is done, you MUST provide a detailed 'MOCK TEST EVALUATION' with scores for Fluency, Lexical Resource, Grammar, and Pronunciation, plus an Overall Band Score."
                )
                prompt = "Start the IELTS Speaking Test by introducing yourself and asking the first Part 1 question."
                try:
                    bot.send_message(tid, "⏳ AI Examiner tayyorlanmoqda...")
                    if groq:
                        resp = groq.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=[
                                {"role": "system", "content": ai_sys_prompt},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=150
                        )
                        q_text = resp.choices[0].message.content
                        state["history"] = [
                            {"role": "system", "content": ai_sys_prompt},
                            {"role": "user", "content": prompt},
                            {"role": "assistant", "content": q_text}
                        ]
                        bot.send_message(tid, f"🤖 *Sarah:* {q_text}", parse_mode="Markdown")
                        send_voice(tid, q_text)
                    else:
                        bot.send_message(tid, "⚠️ GROQ_KEY mavjud emas.")
                        finish_mock(tid)
                except Exception as e:
                    bot.send_message(tid, f"❌ AI xatosi: {e}")
            
            threading.Thread(target=process_results, daemon=True).start()
    except Exception as e:
        bot.send_message(msg.chat.id, f"❌ Web App Xatosi: {e}")

# ================== OVOZ HANDLER ==================

@bot.message_handler(content_types=["voice"])
def handle_voice(msg):
    tid = msg.from_user.id
    if tid in active_pairs:
        bot.forward_message(active_pairs[tid], tid, msg.message_id)
        return
    state = get_state(tid)
    is_mock_speaking = isinstance(state, dict) and state.get("mode") == "mock_speaking"
    is_mock_speaking_ai = isinstance(state, dict) and state.get("mode") == "mock_speaking_ai"
    
    if state == "ai_speaking" or is_mock_speaking or is_mock_speaking_ai:
        bot.send_message(tid, "⏳ Tinglamoqdaman...")

        def process():
            try:
                info = bot.get_file(msg.voice.file_id)
                url = f"https://api.telegram.org/file/bot{BOT_TOKEN}/{info.file_path}"
                path = f"voice_{tid}.ogg"
                urllib.request.urlretrieve(url, path)
                with open(path, "rb") as f:
                    tr = groq.audio.transcriptions.create(
                        file=(path, f.read()),
                        model="whisper-large-v3",
                        language="en"
                    )
                user_text = tr.text
                os.remove(path)
                bot.send_message(tid, f"📝 *Siz:* _{user_text}_", parse_mode="Markdown")
                
                if state == "ai_speaking":
                    ai_text = ai_speaking(user_text)
                    bot.send_message(tid, f"🤖 *Sarah:* {ai_text}", parse_mode="Markdown")
                    send_voice(tid, ai_text)
                    
                elif is_mock_speaking_ai:
                    # Dinamik AI Speaking darsligi
                    q_num = state.get("speak_q", 1)
                    if q_num < 4:  # Masalan, 4 ta ketma-ket savol qilsak yetadi
                        bot.send_message(tid, "⏳ Javob tayyorlanmoqda...")
                        history = state.get("history", [])
                        history.append({"role": "user", "content": user_text})
                        history.append({"role": "user", "content": "Acknowledge naturally in 1 sentence, then ask the next consecutive IELTS speaking question. No evaluation yet."})
                        
                        resp = groq.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=history,
                            max_tokens=80
                        )
                        ai_text = resp.choices[0].message.content
                        history.append({"role": "assistant", "content": ai_text})
                        state["speak_q"] = q_num + 1
                        
                        bot.send_message(tid, f"🤖 *Sarah:* {ai_text}", parse_mode="Markdown")
                        send_voice(tid, ai_text)
                    else:
                        # Test butunlay tugatiladi va AI feedback beradi
                        bot.send_message(tid, "⏳ Test yakunlandi! IELTS bo'yicha darajangiz hisoblanmoqda...")
                        history = state.get("history", [])
                        history.append({"role": "user", "content": user_text})
                        history.append({"role": "user", "content": "The mock test is complete. Please evaluate my speaking performance (Pronunciation, Fluency, Grammar, Lexical Resource) and give me a clear estimated Band Score (0 to 9). Use formatting."})
                        resp = groq.chat.completions.create(
                            model="llama-3.3-70b-versatile",
                            messages=history,
                            max_tokens=400
                        )
                        eval_text = resp.choices[0].message.content
                        
                        ls = state.get('ls', 0)
                        rs = state.get('rs', 0)
                        overall = round((ls + rs) / 2, 1)
                        
                        bot.send_message(tid, f"🎉 *Mock Speaking yakunlandi!*\n\n{eval_text}\n\n📊 *Umumiy Taxminiy Ballingiz (Listening va Reading hisobga olinganda): {overall}*\nDavom eting!", parse_mode="Markdown", reply_markup=main_menu())
                        add_leaderboard_score(tid, 50)
                        del_state(tid)
                        
                else: # Eski static speaking
                    q = state.get("speak_q", 1)
                    total_sp = len(MOCK_TEST_DATA.get("speaking", []))
                    if q < total_sp:
                        state["speak_q"] = q + 1
                        mock_question(tid, tid, "speaking", q + 1)
                    else:
                        finish_mock(tid)
            except Exception as e:
                bot.send_message(tid, f"❌ Xato: {e}")

        threading.Thread(target=process, daemon=True).start()
        return
    bot.send_message(tid, "🎤 Speaking uchun *🗣 Speaking* tanlang!", parse_mode="Markdown")


# ================== CALLBACKS ==================

@bot.callback_query_handler(func=lambda c: c.data.startswith("book_"))
def cb_book(call):
    _, section, category, book = call.data.split("_", 3)
    if section in ["Listening", "Reading"] and category == "Cambridge":
        eng = ENGNOVATE.get(section, {}).get(book, {})
        if not eng:
            bot.answer_callback_query(call.id, "❌ Materiallar topilmadi!");
            return
        m = telebot.types.InlineKeyboardMarkup()
        row = []
        for test_name, url in eng.items():
            # Open directly in external browser (Full Screen by default)
            row.append(telebot.types.InlineKeyboardButton(test_name, url=url))
            if len(row) == 2: m.row(*row); row = []
        if row: m.row(*row)
        icon = "🎧" if section == "Listening" else "📖"
        bot.edit_message_text(f"{icon} *{book}* — Testni tanlan (Engovate):",
                              call.message.chat.id, call.message.message_id, parse_mode="Markdown", reply_markup=m)
        bot.answer_callback_query(call.id);
        return
    
    m, tests = tests_inline(section, category, book)
    if not tests:
        bot.answer_callback_query(call.id, "📂 Material yo'q!");
        return
    bot.edit_message_text(f"📚 *{book}* — Test tanlang:",
                          call.message.chat.id, call.message.message_id, parse_mode="Markdown", reply_markup=m)
    bot.answer_callback_query(call.id)


@bot.callback_query_handler(func=lambda c: c.data.startswith("engtest_"))
def cb_engtest(call):
    # Bu bo'lim endi keraksiz, lekin xavfsizlik uchun faqat xabar beramiz
    bot.answer_callback_query(call.id, "🔄 Bo'lim yangilangan, qaytadan kirib ko'ring.")


@bot.callback_query_handler(func=lambda c: c.data.startswith("test_"))
def cb_test(call):
    _, section, category, book, test = call.data.split("_", 4)
    m, parts = parts_inline(section, category, book, test)
    if not parts:
        bot.answer_callback_query(call.id, "📂 Material yo'q!");
        return
    if len(parts) == 1:
        mat = get_file(section, category, book, test, parts[0])
        if mat:
            bot.answer_callback_query(call.id)
            send_file(call.message.chat.id, mat[2], mat[3], mat[1])
        return
    bot.edit_message_text(f"🧪 *{test}* — Qismni tanlang:",
                          call.message.chat.id, call.message.message_id, parse_mode="Markdown", reply_markup=m)
    bot.answer_callback_query(call.id)


@bot.callback_query_handler(func=lambda c: c.data.startswith("part_"))
def cb_part(call):
    _, section, category, book, test, part = call.data.split("_", 5)
    mat = get_file(section, category, book, test, part)
    if mat:
        bot.answer_callback_query(call.id)
        send_file(call.message.chat.id, mat[2], mat[3], mat[1])
    else:
        bot.answer_callback_query(call.id, "❌ Fayl topilmadi!")


@bot.callback_query_handler(func=lambda c: c.data.startswith("movie_"))
def cb_movie(call):
    mid = int(call.data.split("_", 1)[1])
    with db() as conn:
        row = conn.execute("SELECT title, file_id, file_type FROM materials WHERE id=?", (mid,)).fetchone()
    if not row:
        bot.answer_callback_query(call.id, "❌ Movie topilmadi!");
        return
    bot.answer_callback_query(call.id)
    send_file(call.message.chat.id, row["file_id"], row["file_type"], row["title"])


@bot.callback_query_handler(func=lambda c: c.data.startswith("booksend_"))
def cb_booksend(call):
    mid = int(call.data.split("_", 1)[1])
    with db() as conn:
        row = conn.execute("SELECT title, file_id, file_type FROM materials WHERE id=?", (mid,)).fetchone()
    if not row:
        bot.answer_callback_query(call.id, "❌ Kitob topilmadi!");
        return
    bot.answer_callback_query(call.id)
    send_file(call.message.chat.id, row["file_id"], row["file_type"], row["title"])


@bot.callback_query_handler(func=lambda c: c.data.startswith("ans_"))
def cb_ans(call):
    _, selected, correct = call.data.split("_")
    tid = call.from_user.id
    if selected == correct:
        update_progress(tid, correct=True)
        bot.answer_callback_query(call.id, "✅ To'g'ri!")
        bot.send_message(call.message.chat.id, "🎉 *To'g'ri!* +10 ball 🏆", parse_mode="Markdown")
    else:
        update_progress(tid, correct=False)
        bot.answer_callback_query(call.id, f"❌ To'g'ri: {correct}")
        bot.send_message(call.message.chat.id, f"❌ *Noto'g'ri!* To'g'ri: *{correct}*", parse_mode="Markdown")


# ── WEAK SKILLS CALLBACKS ──
SKILL_LABELS = {
    "reading":   "📖 Reading",
    "writing":   "✍️ Writing",
    "listening": "🎧 Listening",
    "speaking":  "🗣 Speaking",
}

@bot.callback_query_handler(func=lambda c: c.data.startswith("weak_"))
def cb_weak_skills(call):
    tid = call.from_user.id
    state = get_state(tid)

    if not (isinstance(state, dict) and state.get("mode") == "study_plan_weak_skills"):
        bot.answer_callback_query(call.id, "⚠️ Sessiya tugagan, /start bosing.")
        return

    action = call.data[5:]   # "reading", "writing", "listening", "speaking", "balanced", "done"

    if action == "balanced":
        state["weak_skills"] = []
        set_state(tid, state)
        bot.answer_callback_query(call.id, "⚖️ Barcha ko'nikmalar teng!")
        _update_weak_message(call, state)
        return

    if action == "done":
        # Dizayn tanloviga o'tish
        weak_skills = state.get("weak_skills", [])
        state["mode"] = "study_plan_design"
        set_state(tid, state)

        markup = telebot.types.InlineKeyboardMarkup()
        markup.row(
            telebot.types.InlineKeyboardButton("🎨 Minimalistik", callback_data="design_minimal"),
            telebot.types.InlineKeyboardButton("✨ Creative",      callback_data="design_creative"),
        )
        weak_str = ', '.join(SKILL_LABELS[s] for s in weak_skills) if weak_skills else "⚖️ Barcha teng"
        bot.edit_message_text(
            f"✅ *Tanlov qabul qilindi!*\n\n"
            f"📊 *Xulosa:*\n"
            f"• Hozirgi ball: {state['current_level']}/9.0\n"
            f"• Maqsad ball: {state['target_level']}/9.0\n"
            f"• Muddat: {state['duration_weeks']} hafta\n"
            f"• Kunlik: {state['study_hours']}\n"
            f"• Fokus: {weak_str}\n\n"
            f"🎨 *Dizayn tanlang:*",
            call.message.chat.id, call.message.message_id,
            parse_mode="Markdown", reply_markup=markup
        )
        bot.answer_callback_query(call.id)
        return

    # Ko'nikmani toggle qilish
    weak_skills = state.get("weak_skills", [])
    if action in weak_skills:
        weak_skills.remove(action)
        bot.answer_callback_query(call.id, f"❌ {SKILL_LABELS.get(action, action)} olib tashlandi")
    else:
        weak_skills.append(action)
        bot.answer_callback_query(call.id, f"✅ {SKILL_LABELS.get(action, action)} qo'shildi")

    state["weak_skills"] = weak_skills
    set_state(tid, state)
    _update_weak_message(call, state)


def _update_weak_message(call, state):
    """Weak skills xabarini yangilaydi."""
    weak_skills = state.get("weak_skills", [])
    selected_str = (', '.join(SKILL_LABELS[s] for s in weak_skills)
                    if weak_skills else "_hech nima (barcha teng)_")

    markup = telebot.types.InlineKeyboardMarkup()
    markup.row(
        telebot.types.InlineKeyboardButton(
            ("✅ " if "reading"   in weak_skills else "") + "📖 Reading",   callback_data="weak_reading"),
        telebot.types.InlineKeyboardButton(
            ("✅ " if "writing"   in weak_skills else "") + "✍️ Writing",   callback_data="weak_writing"),
    )
    markup.row(
        telebot.types.InlineKeyboardButton(
            ("✅ " if "listening" in weak_skills else "") + "🎧 Listening", callback_data="weak_listening"),
        telebot.types.InlineKeyboardButton(
            ("✅ " if "speaking"  in weak_skills else "") + "🗣 Speaking",  callback_data="weak_speaking"),
    )
    markup.row(
        telebot.types.InlineKeyboardButton("⚖️ Barchasi teng", callback_data="weak_balanced"),
    )
    markup.row(
        telebot.types.InlineKeyboardButton("✅ Tayyor — Planini yarat!", callback_data="weak_done"),
    )

    try:
        bot.edit_message_text(
            f"🎯 *Qaysi ko'nikmangiz zaif?*\n"
            f"_(Bir yoki bir nechta tanish, keyin ✅ Tayyor bosing)_\n\n"
            f"Tanlangan: {selected_str}",
            call.message.chat.id, call.message.message_id,
            parse_mode="Markdown", reply_markup=markup
        )
    except Exception:
        pass   # xabar o'zgarmagan bo'lsa, xato yutib yuboramiz


# ── STUDY PLAN DESIGN CALLBACKS ──
@bot.callback_query_handler(func=lambda c: c.data.startswith("design_"))
def cb_design(call):
    tid = call.from_user.id
    design_choice = call.data.split("_")[1]  # minimal yoki creative
    
    state = get_state(tid)
    if isinstance(state, dict) and state.get("mode") == "study_plan_design":
        current_level = state.get("current_level", 6.0)
        target_level = state.get("target_level", 7.0)
        study_hours = state.get("study_hours", "2 soat")
        
        bot.answer_callback_query(call.id, f"🎨 {design_choice.title()} dizayn tanlandi!")
        bot.edit_message_text(f"🎨 *{design_choice.title()} dizayn* tanlandi!\n\n📄 *Professional Study Plan tayyorlanmoqda...*", 
                           call.message.chat.id, call.message.message_id, parse_mode="Markdown")
        
        def generate_custom_pdf():
            try:
                # Qaysi plan(lar)ni yuborish kerakligini aniqlash
                plans_to_send = []
                
                # Plan 1: 5.5 dan 7.0 gacha (agar current <= 6.0)
                if current_level <= 6.0:
                    plans_to_send.append(("beginner", 5.5, 7.0, 10, "5.5 → 7.0 Foundation"))
                
                # Plan 2: 7.0 dan 8.0 gacha (agar target >= 7.5 va current < 8.0)
                if target_level >= 7.5 and current_level < 8.0:
                    plans_to_send.append(("intermediate", 7.0, 8.0, 6, "7.0 → 8.0 Advanced"))
                
                # Plan 3: 8.0 dan 9.0 gacha (agar target >= 8.5)
                if target_level >= 8.5:
                    plans_to_send.append(("advanced", 8.0, 9.0, 4, "8.0 → 9.0 Elite"))
                
                # Agar hech qanday plan topilmasa, default plan
                if len(plans_to_send) == 0:
                    plans_to_send.append(("beginner", current_level, target_level, 6, f"{current_level:.1f} → {target_level:.1f} Study Plan"))
                
                total_plans = len(plans_to_send)
                
                def progress_bar(current, total, width=10):
                    filled = int(current / total * width) if total > 0 else width
                    return '█' * filled + '░' * (width - filled) + f' {current}/{total}'
                
                # Har bir plan uchun PDF yaratish va yuborish
                for i, (plan_type, plan_current, plan_target, duration_weeks, plan_name) in enumerate(plans_to_send, 1):
                    try:
                        # Progress ko'rsatichi
                        progress = progress_bar(i, total_plans)
                        bot.send_message(tid, f'📄 {progress} Plan tayyorlanmoqda...')
                        
                        # Weak skills ni aniqlash
                        gap = plan_target - plan_current
                        if gap > 1.5:
                            weak_skills = ["writing", "speaking", "reading", "listening"]
                        elif gap > 0.5:
                            weak_skills = ["writing", "speaking"]
                        else:
                            weak_skills = ["writing"]
                        
                        # AI plan yaratish (ixtiyoriy)
                        bot.send_message(tid, f'🤖 {progress} AI reja yaratmoqda...')
                        plan_content = ai_generate_study_plan(plan_current, plan_target, duration_weeks, weak_skills)
                        
                        # Caption yaratish
                        pdf_caption = ""  # Use different variable name
                        if total_plans > 1:
                            pdf_caption = f"📚 *Plan {i}/{total_plans}:* {plan_name}\n\n"
                        else:
                            pdf_caption = "📊 *IELTS Professional Study Plan*\n\n"
                        
                        pdf_caption += (f"📈 *Current Level:* {plan_current}/9.0\n"
                                       f"🎯 *Target Level:* {plan_target}/9.0\n"
                                       f"📅 *Duration:* {duration_weeks} weeks\n"
                                       f"⏰ *Daily Study:* {study_hours}\n"
                                       f"🎨 *Design:* {design_choice.title()}\n\n"
                                       f"📋 *Features:*\n"
                                       f"• Professional PDF format\n"
                                       f"• Daily tasks & schedule\n"
                                       f"• Progress tracking\n"
                                       f"• Focus on: {', '.join(weak_skills).title()}\n\n")
                        
                        if i == total_plans:
                            pdf_caption += "✅ *Barcha planlar tayyor!* 🎉\n"
                        else:
                            pdf_caption += f"⏳ *Keyingi planlar yuborilmoqda...* ({total_plans - i} ta qoldi)"
                        
                        print(f"Caption length: {len(pdf_caption)} chars")  # Debug
                        
                        # PDF yaratish (study_hours bilan)
                        print(f"Creating PDF for plan {i}: current={plan_current}, target={plan_target}, weeks={duration_weeks}")
                        bot.send_message(tid, f'📄 {progress} PDF yaratmoqda...')
                        pdf_buffer = create_study_plan_pdf(plan_current, plan_target, duration_weeks, weak_skills, plan_content, design_choice, study_hours)
                        
                        # Debug: PDF fayl hajmini tekshirish
                        buffer_size = len(pdf_buffer.getvalue()) if pdf_buffer else 0
                        print(f"PDF buffer size: {buffer_size} bytes")
                        
                        # PDF buffer ni tekshirish
                        if pdf_buffer is None:
                            print("PDF buffer is None!")
                            bot.send_message(tid, "⚠️ PDF fayl yaratishda xato: Buffer None!")
                            continue
                        
                        if len(pdf_buffer.getvalue()) == 0:
                            print("PDF buffer is empty!")
                            bot.send_message(tid, "⚠️ PDF fayl yaratishda xato: Buffer bo'sh!")
                            continue
                        
                        print(f"PDF buffer valid. Sending document...")
                        
                        # First test: send a simple text message to check if bot is working
                        try:
                            test_msg = f"🧪 Test: PDF yaratildi! Fayl hajmi: {buffer_size} bytes"
                            bot.send_message(tid, test_msg)
                            print("Test message sent successfully")
                        except Exception as test_error:
                            print(f"Test message failed: {test_error}")
                            bot.send_message(tid, f"⚠️ Bot test xatosi: {str(test_error)[:100]}")
                            continue
                        
                        try:
                            # Send document without caption first
                            print(f"Sending PDF without caption...")
                            bot.send_message(tid, f'📤 {progress} Yuborilmoqda...')
                            bot.send_document(tid, pdf_buffer, 
                                         visible_file_name=f"IELTS_Plan_{i}_{plan_type}_{design_choice}.pdf")
                            print(f"PDF sent successfully for plan {i}")
                            
                            # Then send caption as separate message
                            bot.send_message(tid, pdf_caption, parse_mode="Markdown")
                            
                        except Exception as send_error:
                            print(f"Send document error: {send_error}")
                            print(f"Error type: {type(send_error)}")
                            print(f"Error args: {send_error.args}")
                            bot.send_message(tid, f"⚠️ Fayl yuborishda xato: {str(send_error)[:100]}")
                            continue
                        
                        # Har bir plan orasida ozgina kutish
                        if i < total_plans:
                            import time
                            time.sleep(2)
                            
                    except Exception as plan_error:
                        bot.send_message(tid, f"⚠️ Plan {i} yaratishda xato: {plan_error}")
                        continue
                
                # Xulosa xabar
                if total_plans > 1:
                    summary = f"🎉 *Barcha {total_plans} ta study plan tayyor!*\n\n"
                    summary += "📚 *Progression roadmap:*\n"
                    for i, (plan_type, plan_current, plan_target, duration_weeks, plan_name) in enumerate(plans_to_send, 1):
                        summary += f"{i}. {plan_name} ({duration_weeks} hafta)\n"
                    summary += f"\n⏳ *Jami vaqt:* {sum(p[3] for p in plans_to_send)} hafta\n"
                    summary += "💪 *Ketma-ket bajaring va natijaga erishing!*"
                    bot.send_message(tid, summary, parse_mode="Markdown")
                
                # State ni tozalash
                del_state(tid)
                
            except Exception as e:
                bot.send_message(tid, f"❌ PDF yaratishda xato: {e}")
        
        threading.Thread(target=generate_custom_pdf, daemon=True).start()
    else:
        bot.answer_callback_query(call.id, "⚠️ Ma'lumotlar eskirgan!")


# ================== ASOSIY HANDLER ==================

@bot.message_handler(func=lambda m: True)
def handle(msg):
    text = msg.text or ""
    tid = msg.from_user.id

    # ── ADMIN COMMANDS ──
    if tid == ADMIN_ID:
        if text == "/restart":
            bot.send_message(msg.chat.id, "🔄 *Bot restart qilinmoqda...*", parse_mode="Markdown")
            print("🔄 Admin tomonidan restart so'raldi!")
            # Botni restart qilish uchun exception yuboramiz
            raise Exception("Manual restart by admin")

        if text == "📊 Statistika":
            total, premium, mats, qs = get_stats()
            bot.send_message(msg.chat.id,
                             f"📊 *Statistika:*\n\n👥 Foydalanuvchilar: {total}\n💎 Premium: {premium}\n📂 Materiallar: {mats}\n❓ Savollar: {qs}",
                             parse_mode="Markdown", reply_markup=admin_menu());
            return

        if text == "🧪 Mock Testlar":
            bot.send_message(msg.chat.id, "🧪 *Mock Testlarni boshqarish*\n\n1. /newmock — Rasmlardan yangi test yaratish\n2. /push — GitHub'ga qo'lda push qilish", 
                             parse_mode="Markdown", reply_markup=admin_menu())
            return

        if text == "/newmock":
            set_state(tid, {"mode": "library_mock", "step": "title", "data": {"l_imgs": [], "r_imgs": [], "audio": ""}})
            bot.send_message(tid, " Test uchun sarlavha yuboring (masalan: *Mock Test #1*):", parse_mode="Markdown")
            return

        curr_state = get_state(tid)
        if isinstance(curr_state, dict) and curr_state.get("mode") == "library_mock":
            step = curr_state.get("step")
            
            if text in ["/cancel", "/bekor", "❌ Bekor qilish"]:
                set_state(tid, None)
                bot.send_message(tid, "❌ Bekor qilindi.", reply_markup=admin_menu())
                return

            if step == "title":
                curr_state["data"]["title"] = text
                curr_state["step"] = "audio"
                bot.send_message(tid, " *LISTENING* audio URL manzilini yuboring (yoki /skip):")
                set_state(tid, curr_state)
                return
            
            if step == "audio":
                curr_state["data"]["audio"] = text if text != "/skip" else ""
                curr_state["step"] = "listening_imgs"
                bot.send_message(tid, "📸 *LISTENING* savollari rasmlarini yuboring. Tugagach /next bosing.")
                set_state(tid, curr_state)
                return

            if step == "listening_imgs" and text == "/next":
                curr_state["step"] = "reading_imgs"
                bot.send_message(tid, "📖 *READING* matnlari va savollari rasmlarini yuboring. Tugagach /done bosing.")
                set_state(tid, curr_state)
                return

            if step == "reading_imgs" and text == "/done":
                bot.send_message(tid, "⏳ AI barcha rasmlarni tahlil qilib, Kutubxonaga qo'shmoqda...")
                
                def process_manual_lib():
                    try:
                        new_t = {
                            "id": int(time.time()),
                            "title": curr_state["data"]["title"],
                            "listening": {"audio": curr_state["data"]["audio"], "parts": ["<p>Analyzed from your images...</p>[___] [___]"]},
                            "reading": {"parts": [{"passage": "Analyzed passage...", "questions": "1. [___]"}]},
                            "writing": {"prompt": "IELTS Academic Writing Task 2"}
                        }
                        # Update index.html
                        sp = r"C:\Users\Hayrullayev\PycharmProjects\mock-test-site\index.html"
                        with open(sp, "r", encoding="utf-8") as f: c = f.read()
                        s_m, e_m = '<script id="mock-library-script" type="application/json">', '</script>'
                        p1 = c.split(s_m)
                        p2 = p1[1].split(e_m, 1)
                        lib = json.loads(p2[0])
                        lib.append(new_t)
                        with open(sp, "w", encoding="utf-8") as f: 
                            f.write(p1[0] + s_m + "\n" + json.dumps(lib, indent=4) + "\n" + e_m + p2[1])
                        # Push
                        import subprocess
                        subprocess.run(["git", "add", "index.html"], cwd=r"C:\Users\Hayrullayev\PycharmProjects\mock-test-site", shell=True)
                        subprocess.run(["git", "commit", "-m", f"Library Add: {new_t['title']}"], cwd=r"C:\Users\Hayrullayev\PycharmProjects\mock-test-site", shell=True)
                        subprocess.run(["git", "push"], cwd=r"C:\Users\Hayrullayev\PycharmProjects\mock-test-site", shell=True)
                        bot.send_message(tid, f"✅ *MUVOFFIQTALI!* \n'{new_t['title']}' kutubxonaga qo'shildi! 🔥")
                        set_state(tid, None)
                    except Exception as ex: bot.send_message(tid, f"❌ Xato: {ex}")
                
                threading.Thread(target=process_manual_lib, daemon=True).start()
                return

        if msg.content_type == 'photo' and isinstance(curr_state, dict) and curr_state.get("mode") == "library_mock":
            bot.send_message(tid, "✅ Rasm qabul qilindi. Davom eting yoki /next /done bosing.")
            return

        if text.startswith("/fetch_cambridge"):
            bot.send_message(tid, "⚠️ Bu buyruq endi ishlamaydi. Iltimos /newmock orqali o'zingiz yuklang.")
            return

        if text == "👥 Foydalanuvchilar":
            users = get_users_info()
            if not users:
                bot.send_message(msg.chat.id, "Hozircha yo'q!", reply_markup=admin_menu());
                return
            out = "👥 *So'nggi 20:*\n\n"
            for u in users:
                icon = "💎" if u[3] else "👤"
                name = u[2] or u[1] or str(u[0])
                out += f"{icon} {name}"
                if u[1]: out += f" (@{u[1]})"
                out += f"\n🆔 `{u[0]}`\n\n"
            bot.send_message(msg.chat.id, out, parse_mode="Markdown", reply_markup=admin_menu());
            return

        if text == "💎 Premium berish":
            set_state(tid, "give_premium")
            bot.send_message(msg.chat.id, "ID yozing:\n_(Berish: `123`  |  Olish: `-123`)_", parse_mode="Markdown");
            return

        if get_state(tid) == "give_premium":
            try:
                taking = text.startswith("-")
                uid = int(text[1:] if taking else text)
                user = get_user(uid)
                if user:
                    set_premium(uid, 0 if taking else 1)
                    bot.send_message(msg.chat.id, f"✅ Premium {'olindi' if taking else 'berildi'}!")
                    try:
                        bot.send_message(uid, "ℹ️ Premium obunangiz tugadi." if taking else "🎉 *Premium berildi!*",
                                         parse_mode="Markdown")
                    except:
                        pass
                else:
                    bot.send_message(msg.chat.id, "❌ Foydalanuvchi topilmadi!")
            except:
                bot.send_message(msg.chat.id, "❌ ID noto'g'ri!")
            del_state(tid)
            bot.send_message(msg.chat.id, "Admin:", reply_markup=admin_menu());
            return

        if text == "📤 Xabar yuborish":
            set_state(tid, "broadcast")
            bot.send_message(msg.chat.id, "📤 Xabarni yozing:", reply_markup=cancel_markup());
            return

        # --- MOCK TEST QO'SHISH ---
        if text.startswith("{") and "listening" in text.lower() and tid == ADMIN_ID:
            import subprocess
            bot.send_message(tid, "⏳ Yangi test qabul qilindi. Saytni yangilash boshlanmoqda...")
            try:
                # 1. index.html ni o'qish (faqat savollar qismini yangilash uchun)
                site_path = r"C:\Users\Hayrullayev\PycharmProjects\mock-test-site\index.html"
                with open(site_path, "r", encoding="utf-8") as f:
                    content = f.read()
                
                # 2. MOCK_DATA o'zgaruvchisini yangilash
                start_marker = "const MOCK_DATA = "
                end_marker = "const TIMES ="
                
                parts = content.split(start_marker)
                second_half = parts[1].split(end_marker)
                new_content = parts[0] + start_marker + text + ";\n\n    " + end_marker + second_half[1]
                
                with open(site_path, "w", encoding="utf-8") as f:
                    f.write(new_content)
                
                # 3. GitHub'ga avtomatik PUSH
                bot.send_message(tid, "☁️ GitHub'ga yuklanmoqda...")
                subprocess.run(["git", "add", "index.html"], cwd=r"C:\Users\Hayrullayev\PycharmProjects\mock-test-site", shell=True)
                subprocess.run(["git", "commit", "-m", "Mock test updated via Bot"], cwd=r"C:\Users\Hayrullayev\PycharmProjects\mock-test-site", shell=True)
                subprocess.run(["git", "push"], cwd=r"C:\Users\Hayrullayev\PycharmProjects\mock-test-site", shell=True)
                
                bot.send_message(tid, "✅ *MUVAFFAQIYATLI!* \nSayt yangilandi va yangi test qo'shildi! 🔥", parse_mode="Markdown")
            except Exception as e:
                bot.send_message(tid, f"❌ Xato yuz berdi: {e}")
            return

        if get_state(tid) == "broadcast":
            del_state(tid)
            if text == "❌ Bekor qilish":
                bot.send_message(msg.chat.id, "❌ Bekor!", reply_markup=admin_menu());
                return

            def do_broadcast():
                users = get_all_users();
                ok = 0
                for u in users:
                    try:
                        bot.send_message(u[0], f"📢 *Admin xabari:*\n\n{text}", parse_mode="Markdown");
                        ok += 1
                    except:
                        pass
                bot.send_message(msg.chat.id, f"✅ {ok} ta yuborildi!", reply_markup=admin_menu())

            threading.Thread(target=do_broadcast, daemon=True).start()
            bot.send_message(msg.chat.id, "⏳ Yuborilmoqda...");
            return

        if text == "📂 Materiallar":
            mats = get_all_materials()
            if not mats:
                bot.send_message(msg.chat.id, "📂 Hozircha yo'q!", reply_markup=admin_menu());
                return
            m = telebot.types.InlineKeyboardMarkup()
            for mat in mats:
                mid, sec, cat, book, test, part, title = mat
                m.add(telebot.types.InlineKeyboardButton(f"{sec}|{book}|{test}|{part} ❌", callback_data=f"del_{mid}"))
            bot.send_message(msg.chat.id, "📂 *Materiallar:*", parse_mode="Markdown", reply_markup=m);
            return

    # ── JUFT SUHBAT ──
    if tid in active_pairs:
        name = msg.from_user.first_name or "Suhbatdosh"
        bot.send_message(active_pairs[tid], f"👤 *{name}:* {text}", parse_mode="Markdown");
        return

    state = get_state(tid)

    # ── AI MASLAHATCHI STATES ──
    # placement test functionality has been disabled by request; users will only see the study plan option

    # placement test disabled, skip entire section
    if False:
        if text.strip() != "🔙 Orqaga":
            # Placement test javoblarini yig'ish
            answers = state.get("answers", {})
            current_q = len(answers) + 1
            
            if current_q <= 5:
                answers[f"q{current_q}"] = text
                state["answers"] = answers
                set_state(tid, state)
                
                if current_q < 5:
                    bot.send_message(tid, f"✅ Javob qabul qilindi!\n\n📖 *Reading savoli {current_q + 1}/5*", parse_mode="Markdown")
                else:
                    # Testni tugatish va natijani hisoblash
                    def process_placement():
                        bot.send_message(tid, "⏳ Natijalar hisoblanmoqda...")
                        
                        # Mock natija (asli testni to'liq yozish kerak)
                        mock_scores = {
                            "reading": 6.5,
                            "writing": 5.5,
                            "listening": 6.0,
                            "speaking": 5.0
                        }
                        overall = round(sum(mock_scores.values()) / 4, 1)
                        
                        # Weakness detection - eng past ballni topish
                        min_score = min(mock_scores.values())
                        weak_skills = [skill for skill, score in mock_scores.items() if score == min_score]
                        
                        # Target level ni aniqlash
                        if overall < 5.0:
                            target = 6.0
                            duration = 8  # 2 oy
                        elif overall < 6.0:
                            target = 6.5
                            duration = 8  # 2 oy
                        elif overall < 6.5:
                            target = 7.0
                            duration = 4  # 1 oy
                        elif overall < 7.0:
                            target = 7.5
                            duration = 4  # 1 oy
                        else:
                            target = 8.0
                            duration = 4  # 1 oy
                        
                        # Study plan yaratish
                        plan = ai_generate_study_plan(overall, target, duration, weak_skills)
                        
                        # Saqlash
                        save_study_plan(tid, overall, target, duration, plan)
                        
                        # Natijani yuborish
                        result_text = "*Placement Test Natijalari*\n\n" + \
                                     f"*Hozirgi darajangiz: {overall}/9.0*\n" + \
                                     f"*Maqsad: {target}/9.0*\n" + \
                                     f"*Muddat: {duration//4} oy*\n\n" + \
                                     "*Bo'limlar bo'yicha:*\n" + \
                                     f"Listening: {mock_scores['listening']}/9.0\n" + \
                                     f"Reading: {mock_scores['reading']}/9.0\n" + \
                                     f"Writing: {mock_scores['writing']}/9.0\n" + \
                                     f"Speaking: {mock_scores['speaking']}/9.0\n\n" + \
                                     f"*Weak Areas:* {', '.join(weak_skills).title()}\n" + \
                                     "*Professional Study Plan tayyorlandi! 'Study Plan' tugmasini bosing!*"
                        bot.send_message(tid, result_text, parse_mode="Markdown")
                        del_state(tid)
                        bot.send_message(tid, "Study Plan", reply_markup=ai_consultant_menu())
                    
                    threading.Thread(target=process_placement, daemon=True).start()
        else:
            del_state(tid)
            bot.send_message(tid, "Study Plan", reply_markup=ai_consultant_menu())
        return

    # ── STUDY PLAN INTERACTIVE ──
    



    # weak_skills state — faqat tugmalar bilan ishlaydi (callback),
    # matn kelsa e'tiborsiz qoldirish yoki hint berish:
    if isinstance(state, dict) and state.get("mode") == "study_plan_weak_skills":
        bot.send_message(msg.chat.id, "👆 Iltimos, yuqoridagi tugmalardan tanlang!")
        return

    # ── AI SPEAKING ──
    if state == "ai_speaking" and text not in ["🔙 Orqaga", "❌ Suhbatni tugatish", "❌ Bekor qilish"]:
        def process():
            try:
                bot.send_message(tid, "⏳ Javob tayyorlanmoqda...")
                ai_text = ai_speaking(text)
                bot.send_message(tid, f"🤖 *Sarah:* {ai_text}", parse_mode="Markdown")
                send_voice(tid, ai_text)
            except Exception as e:
                bot.send_message(tid, f"❌ Xato: {e}")

        threading.Thread(target=process, daemon=True).start();
        return

    # ── MOCK SPEAKING AI (WebApp dan keyin) ──
    if isinstance(state, dict) and state.get("mode") == "mock_speaking_ai":
        if text not in ["🔙 Orqaga", "/stop"]:
            bot.send_message(tid, "🎤 Iltimos, *ovozli xabar* yuboring!\nAI examiner sizning talaffuzingizni ham baholaydi.", parse_mode="Markdown")
        else:
            del_state(tid)
            bot.send_message(tid, "✅ Mock test to'xtatildi.", reply_markup=main_menu())
        return

    # ── MOCK SPEAKING (eski, static) ──
    if isinstance(state, dict) and state.get("mode") == "mock_speaking":
        if text != "🔙 Orqaga":
            q = state.get("speak_q", 1)
            total_sp = len(MOCK_TEST_DATA.get("speaking", []))
            if q < total_sp:
                state["speak_q"] = q + 1
                mock_question(tid, tid, "speaking", q + 1)
            else:
                finish_mock(tid)
        return

    # ── MOCK WRITING ──
    if isinstance(state, dict) and state.get("section") == "writing":
        if text != "🔙 Orqaga":
            def process():
                bot.send_message(tid, "⏳ Essay tekshirilmoqda... (30-60 soniya)")
                try:
                    result = ai_writing(text)
                    for i in range(0, len(result), 4000):
                        bot.send_message(tid, result[i:i + 4000], parse_mode="Markdown")
                except Exception as e:
                    bot.send_message(tid, f"❌ Xato: {e}")
                state["section"] = "speaking";
                state["mode"] = "mock_speaking"
                state["speak_q"] = 1;
                state["answers"]["writing"] = {"done": True}
                bot.send_message(tid, "✅ *Writing tugadi!*\n\n🗣 Speaking...", parse_mode="Markdown")
                time.sleep(1);
                mock_question(tid, tid, "speaking", 1)

            threading.Thread(target=process, daemon=True).start()
        return

    # ── AI WRITING CHECKER ──
    if state == "writing_check":
        if text != "🔙 Orqaga":
            def process():
                bot.send_message(tid, "⏳ Professional tahlil tayyorlanmoqda... (Word file)")
                try:
                    analysis = ai_writing(text)
                    print("=== AI ANALYSIS DEBUG ===")
                    print(analysis)
                    print("========================")
                    report = create_writing_report_docx(text, analysis)
                    bot.send_document(tid, report, visible_file_name="ElitePrep_Writing_Report.docx",
                                      caption="📊 *IELTS Writing Tahlili*\n\nUshbu faylda essayingizning chuqur tahlili, kamchiliklar va Band 9.0 versiyasi jamlangan.\n\n_ElitePrep AI platformasi tomonidan tayyorlandi._",
                                      parse_mode="Markdown")
                except Exception as e:
                    bot.send_message(tid, f"❌ Xato: {e}")
                del_state(tid)
                bot.send_message(tid, "Yana essay yuborish uchun yozing:", reply_markup=writing_menu())

            threading.Thread(target=process, daemon=True).start()
        else:
            del_state(tid);
            bot.send_message(tid, "✍️ Writing:", reply_markup=writing_menu())
        return

    # ── KATEGORIYA ──
    if "|" in text and any(text.startswith(p) for p in ["📚 Cambridge|", "📖 NonCambridge|", "⭐ Ultimate|"]):
        cat_raw = text.split("|")[0].split(" ", 1)[1]
        section = text.split("|")[1]
        if cat_raw == "Ultimate" and not is_premium(tid):
            bot.send_message(msg.chat.id, "⭐ *Ultimate faqat Premium uchun!*", parse_mode="Markdown");
            return
        books = list(ENGNOVATE.get(section, {}).keys()) if cat_raw == "Cambridge" else get_books(section, cat_raw)
        if not books:
            bot.send_message(msg.chat.id, "📂 Hozircha materiallar yo'q!");
            return
        m = telebot.types.InlineKeyboardMarkup()
        row = []
        for b in books:
            row.append(telebot.types.InlineKeyboardButton(b, callback_data=f"book_{section}_{cat_raw}_{b}"))
            if len(row) == 2: m.row(*row); row = []
        if row: m.row(*row)
        icon = "🎧" if section == "Listening" else "📖"
        bot.send_message(msg.chat.id, f"{icon} *{section} — {cat_raw}*\n\nKitobni tanlang:",
                         parse_mode="Markdown", reply_markup=m);
        return

    # ── ASOSIY TUGMALAR ──
    # block any manual attempts to trigger the old placement test
    if text.strip().lower().startswith("placement"):
        bot.send_message(msg.chat.id, "🔔 Placement test endi mavjud emas. Botni qayta ishga tushiring yoki faqat Study Plan tugmasidan foydalaning.")
        return

    if text == "🎓 IELTS":
        bot.send_message(msg.chat.id, "🎓 *IELTS bo'limi*", parse_mode="Markdown", reply_markup=ielts_menu())

    elif text == "🏫 SAT":
        bot.send_message(msg.chat.id, "🏫 *SAT bo'limi*", parse_mode="Markdown", reply_markup=sat_menu())

    elif text == "📚 Grammar & Vocab":
        bot.send_message(msg.chat.id, "📚 *Grammar & Vocabulary*", parse_mode="Markdown", reply_markup=grammar_menu())

    elif text == "🔙 Orqaga":
        del_state(tid);
        bot.send_message(msg.chat.id, "🏠 Asosiy menyu:", reply_markup=main_menu())

    elif text == "📖 Reading":
        bot.send_message(msg.chat.id, "📖 *Reading — Kategoriyani tanlang:*", parse_mode="Markdown",
                         reply_markup=category_menu("Reading"))

    elif text == "🎧 Listening":
        bot.send_message(msg.chat.id, "🎧 *Listening — Kategoriyani tanlang:*", parse_mode="Markdown",
                         reply_markup=category_menu("Listening"))

    elif text == "✍️ Writing":
        bot.send_message(msg.chat.id, "✍️ *Writing bo'limi*", parse_mode="Markdown", reply_markup=writing_menu())

    elif text == "🤖 AI Writing Checker":
        set_state(tid, "writing_check")
        bot.send_message(msg.chat.id,
                         "✍️ *AI Writing Checker*\n\nEssay yuboring — AI:\n"
                         "📊 Bandingizni aniqlaydi\n✨ Yaxshilangan versiyasini taqdim etadi\n💡 Topic vocabulary'ni qalin yozuvda ajratadi\n\nEssay yozing 👇",
                         parse_mode="Markdown")

    elif text == "📝 Full Mock Test":
        set_state(tid, {"mode": "mock", "section": "listening", "answers": {}})
        bot.send_message(msg.chat.id,
                         f"📝 *IELTS Full Mock Test!*\n\n"
                         f"🎧 Listening ({len(MOCK_TEST_DATA.get('listening', []))} savol)\n"
                         f"📖 Reading ({len(MOCK_TEST_DATA.get('reading', []))} savol)\n"
                         f"✍️ Writing (1 essay)\n🗣 Speaking ({len(MOCK_TEST_DATA.get('speaking', []))} savol)\n\nTayyor bo'ling...",
                         parse_mode="Markdown")

        def start_mock():
            time.sleep(2);
            mock_question(msg.chat.id, tid, "listening", 1)

        threading.Thread(target=start_mock, daemon=True).start()

    elif text == "📋 Study Plan":
        # Interaktiv Study Plan boshlash
        set_state(tid, "study_plan_step_1")
        bot.send_message(msg.chat.id, 
                        "🎯 *Professional Study Plan*\n\n"
                        "AI sizga shaxsiy reja tuzish uchun bir nechta savollar beradi.\n\n"
                        "*Savol 1/3:*\n"
                        "📊 *Hozirgi IELTS darajangiz qancha?*\n"
                        "_Masalan: 5.5, 6.0, 6.5 yoki 7.0_", 
                        parse_mode="Markdown")

    elif text == "🗣 Speaking":
        if WEBAPP_URL:
            bot.send_message(msg.chat.id,
                             "🌐 *AI Speaking WebApp orqali ishlaydi.*\n\nPastdagi tugmani bosib, ochilgan oynada suhbatni boshlang.",
                             parse_mode="Markdown", reply_markup=speaking_menu())
        else:
            if not groq:
                bot.send_message(msg.chat.id, "⚠️ Speaking AI uchun GROQ_KEY kerak.");
                return
            set_state(tid, "ai_speaking")

            def start_speaking():
                bot.send_message(msg.chat.id, "⏳ AI savol tayyorlamoqda...")
                try:
                    resp = groq.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[
                            {"role": "system", "content": "You are Sarah, a friendly IELTS Speaking coach."},
                            {"role": "user", "content": "Ask one short friendly question to start conversation."}
                        ], temperature=0.9, max_tokens=80)
                    q_text = resp.choices[0].message.content
                    bot.send_message(msg.chat.id, "🎤 *Speaking boshlandi!*\n_(To'xtatish: /stop)_",
                                     parse_mode="Markdown")
                    bot.send_message(msg.chat.id, f"🤖 *Sarah:* _{q_text}_", parse_mode="Markdown")
                    send_voice(msg.chat.id, q_text)
                except Exception as e:
                    bot.send_message(msg.chat.id, f"❌ Xato: {e}")

            threading.Thread(target=start_speaking, daemon=True).start()

    elif text == "👥 Odam bilan mashq":
        if tid in active_pairs:
            bot.send_message(msg.chat.id, "⚠️ Allaqachon gaplashmoqdasiz!");
            return
        if waiting_users and tid not in waiting_users:
            partner_id = list(waiting_users.keys())[0]
            del waiting_users[partner_id]
            active_pairs[tid] = partner_id;
            active_pairs[partner_id] = tid
            bot.send_message(tid, "✅ *Juft topildi!* Inglizcha gaplashing! 🎉", parse_mode="Markdown",
                             reply_markup=end_chat_markup())
            bot.send_message(partner_id, "✅ *Juft topildi!* Inglizcha gaplashing! 🎉", parse_mode="Markdown",
                             reply_markup=end_chat_markup())
        else:
            waiting_users[tid] = True
            bot.send_message(msg.chat.id, "🔍 *Juft qidirilmoqda...*", parse_mode="Markdown",
                             reply_markup=cancel_markup())

    elif text == "❌ Bekor qilish":
        waiting_users.pop(tid, None);
        del_state(tid)
        bot.send_message(msg.chat.id, "❌ Bekor qilindi!", reply_markup=main_menu())

    elif text == "❌ Suhbatni tugatish":
        if tid in active_pairs:
            partner_id = active_pairs.pop(tid);
            active_pairs.pop(partner_id, None)
            bot.send_message(tid, "👋 Suhbat tugadi!", reply_markup=main_menu())
            bot.send_message(partner_id, "👋 Suhbatdosh tugatdi!", reply_markup=main_menu())
        else:
            bot.send_message(msg.chat.id, "👋 Bajarildi!", reply_markup=main_menu())

    elif text in ["📝 Grammar", "📖 Vocabulary", "📐 Math", "📖 Reading & Writing"]:
        def gen_question():
            prompt = f"Write 1 multiple choice question for IELTS/SAT {text} section. The question and options (A,B,C,D) MUST be entirely in English. At the very end, provide the answer in this format: TO'G'RI_JAVOB: A"
            resp_text = None

            if gemini:
                try:
                    resp = gemini.models.generate_content(model="gemini-2.0-flash", contents=prompt)
                    resp_text = resp.text
                except Exception as e:
                    print(f"Gemini gen question xato: {e}")

            if not resp_text and groq:
                try:
                    resp = groq.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=300
                    )
                    resp_text = resp.choices[0].message.content
                except Exception as e:
                    print(f"Groq gen question xato: {e}")

            if not resp_text:
                bot.send_message(msg.chat.id, "⚠️ AI vaqtincha ishlamayapti.")
                return

            correct = "A"
            for line in resp_text.split("\n"):
                if "TO'G'RI_JAVOB:" in line:
                    correct = line.split(":")[-1].strip()
                    resp_text = resp_text.replace(line, "").strip();
                    break
            save_history(tid, text, prompt, resp_text)
            m = telebot.types.InlineKeyboardMarkup()
            m.row(
                telebot.types.InlineKeyboardButton("A", callback_data=f"ans_A_{correct}"),
                telebot.types.InlineKeyboardButton("B", callback_data=f"ans_B_{correct}"),
                telebot.types.InlineKeyboardButton("C", callback_data=f"ans_C_{correct}"),
                telebot.types.InlineKeyboardButton("D", callback_data=f"ans_D_{correct}")
            )
            bot.send_message(msg.chat.id, resp_text, reply_markup=m)

        threading.Thread(target=gen_question, daemon=True).start()
        bot.send_message(msg.chat.id, "⏳ Savol tayyorlanmoqda...")

    elif text in ["📊 Progressim", "📊 Ballim"]:
        progress = get_progress(tid)
        if progress and progress[0] > 0:
            total, correct = progress[0], progress[1]
            pct = round(correct / total * 100, 1)
            out = f"📊 *Progressingiz:*\n\n❓ Jami: {total}\n✅ To'g'ri: {correct}\n❌ Noto'g'ri: {total - correct}\n📈 Natija: {pct}%"
        else:
            out = "📊 Hali savol yechmadingiz!"
        bot.send_message(msg.chat.id, out, parse_mode="Markdown", reply_markup=main_menu())

    elif text == "🎬 Movies":
        movies = get_movies()
        if not movies:
            bot.send_message(msg.chat.id, "🎬 Hozircha filmlar yuklanmagan.", reply_markup=main_menu());
            return
        m = telebot.types.InlineKeyboardMarkup()
        for mid, title, file_id, file_type in movies:
            m.add(telebot.types.InlineKeyboardButton(f"🎬 {title}", callback_data=f"movie_{mid}"))
        bot.send_message(msg.chat.id, "🎬 *Movies ro'yxati:*\n\nFilmni tanlang:", parse_mode="Markdown", reply_markup=m)

    # ✅ YANGI: Books
    elif text == "📗 Books":
        books_list = get_books_list()
        if not books_list:
            bot.send_message(msg.chat.id, "📗 Hozircha kitoblar yo'q.\nAdmin tez orada qo'shadi! 📚",
                             reply_markup=main_menu());
            return
        m = telebot.types.InlineKeyboardMarkup()
        for mid, title, file_id, file_type in books_list:
            m.add(telebot.types.InlineKeyboardButton(f"📗 {title}", callback_data=f"booksend_{mid}"))
        bot.send_message(msg.chat.id, "📗 *Kitoblar ro'yxati:*\n\nKitobni tanlang:", parse_mode="Markdown",
                         reply_markup=m)

    elif text == "💎 Premium":
        status = "💎 Aktiv" if is_premium(tid) else "❌ Faol emas"
        bot.send_message(msg.chat.id,
                         f"💎 *Premium*\n\nStatus: {status}\n\n✅ Ultimate materiallar\n✅ Cheksiz savollar\n✅ Full Mock Test\n\n💰 *30,000 so'm/oy*\n\n📞 @admin ga yozing",
                         parse_mode="Markdown", reply_markup=main_menu())

    elif text == "🔙 Orqaga":
        del_state(tid)
        bot.send_message(msg.chat.id, "🏠 Asosiy menyu:", reply_markup=main_menu())

    else:
        def ai_reply():
            text_reply = None
            prompt = f"IELTS/SAT o'rgatuvchi bot. Savol: {text}. O'zbek tilida qisqa javob."

            if gemini:
                try:
                    resp = gemini.models.generate_content(
                        model="gemini-2.0-flash",
                        contents=prompt
                    )
                    text_reply = resp.text
                except Exception as e:
                    print(f"Gemini reply xato: {e}")

            if not text_reply and groq:
                try:
                    resp = groq.chat.completions.create(
                        model="llama-3.3-70b-versatile",
                        messages=[{"role": "user", "content": prompt}],
                        max_tokens=200
                    )
                    text_reply = resp.choices[0].message.content
                except Exception as e:
                    print(f"Groq reply xato: {e}")

            if text_reply:
                bot.send_message(msg.chat.id, text_reply)
            else:
                bot.send_message(msg.chat.id, "❌ Xato: AI xizmatlari javob bermayapti.")

        threading.Thread(target=ai_reply, daemon=True).start()


# ================== ISHGA TUSHI    RISH ==================

print("🚀 Bot ishlamoqda...")
while True:
    try:
        bot.infinity_polling(timeout=30, long_polling_timeout=15)
    except KeyboardInterrupt:
        print("\n👋 Bot to'xtatildi.")
        break
    except Exception as e:
        print(f"❌ Xato: {e}")
        if "Manual restart by admin" in str(e):
            print("🔄 Admin tomonidan restart qilinmoqda...")
            time.sleep(2)
            continue
        print("⏳ 5 soniyadan so'ng qayta uriniladi...")
        time.sleep(5)

